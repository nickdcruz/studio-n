import asyncio
import json
import logging
import os
import re
import secrets
import sqlite3
import traceback
import uuid
from datetime import datetime
from pathlib import Path

import uvicorn

import httpx
import requests
from apify_client import ApifyClient
from anthropic import AsyncAnthropic
from dotenv import load_dotenv
from fastapi import FastAPI, Request, UploadFile, File, Form
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse, StreamingResponse
from typing import List
from fastapi.templating import Jinja2Templates
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.middleware.sessions import SessionMiddleware
from starlette.responses import Response

load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
    force=True,
)

BASE_DIR    = Path(__file__).parent
OUTPUTS_DIR = BASE_DIR / "outputs"
OUTPUTS_DIR.mkdir(exist_ok=True)

# ── Persistent storage (SQLite) ───────────────────────────────────
# Set DATA_DIR=/data in Railway env vars + mount a Railway Volume at /data.
# Without a volume, data persists in ./data (lost on redeploy — set the volume).
DATA_DIR = Path(os.getenv("DATA_DIR", BASE_DIR / "data"))
DATA_DIR.mkdir(exist_ok=True)
DB_PATH = DATA_DIR / "studio_n.db"

def init_db():
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS jobs (
            id              TEXT PRIMARY KEY,
            title           TEXT NOT NULL DEFAULT '',
            client          TEXT NOT NULL DEFAULT 'other',
            project_name    TEXT DEFAULT '',
            brief           TEXT NOT NULL DEFAULT '',
            timestamp       TEXT NOT NULL,
            marcus_analysis TEXT DEFAULT '',
            stage1_outputs  TEXT DEFAULT '{}',
            marcus_review   TEXT DEFAULT '',
            cascade_outputs TEXT DEFAULT '{}',
            marcus_cascade_review TEXT DEFAULT '',
            assembled_files TEXT DEFAULT '{}',
            assembled_content TEXT DEFAULT '{}',
            status          TEXT DEFAULT 'complete',
            created_at      TEXT DEFAULT (datetime('now'))
        )
    """)
    try:
        conn.execute("ALTER TABLE jobs ADD COLUMN archived INTEGER DEFAULT 0")
        conn.commit()
    except Exception:
        pass
    conn.commit()
    conn.close()
    logging.info("DB ready at %s", DB_PATH)

def db_save_job(job_id: str, job: dict, assembled_content: dict = None):
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        INSERT OR REPLACE INTO jobs
        (id, title, client, project_name, brief, timestamp,
         marcus_analysis, stage1_outputs, marcus_review,
         cascade_outputs, marcus_cascade_review,
         assembled_files, assembled_content, status)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)
    """, (
        job_id,
        job.get("title", "Untitled"),
        job.get("client", "other"),
        job.get("project_name", ""),
        job.get("brief", ""),
        job["timestamp"].isoformat() if hasattr(job.get("timestamp"), "isoformat") else str(job.get("timestamp", "")),
        job.get("marcus_analysis", ""),
        json.dumps(job.get("stage1_outputs") or {}),
        job.get("marcus_review", "") or "",
        json.dumps(job.get("cascade_outputs") or {}),
        job.get("marcus_cascade_review", "") or "",
        json.dumps(job.get("assembled_files") or {}),
        json.dumps(assembled_content or {}),
        "complete",
    ))
    conn.commit()
    conn.close()

def db_save_job_initial(job_id: str, title: str, client: str, project_name: str, brief: str):
    """Write a processing stub so the job appears in Outputs immediately."""
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        INSERT OR IGNORE INTO jobs
        (id, title, client, project_name, brief, timestamp, status)
        VALUES (?,?,?,?,?,datetime('now'),'processing')
    """, (job_id, title or brief[:60], client, project_name, brief))
    conn.commit()
    conn.close()

def db_update_job_status(job_id: str, status: str):
    conn = sqlite3.connect(DB_PATH)
    conn.execute("UPDATE jobs SET status=? WHERE id=?", (status, job_id))
    conn.commit()
    conn.close()

def db_load_all_jobs(show_archived: bool = False) -> list[dict]:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    if show_archived:
        rows = conn.execute("SELECT * FROM jobs WHERE archived=1 ORDER BY created_at DESC").fetchall()
    else:
        rows = conn.execute("SELECT * FROM jobs WHERE archived=0 OR archived IS NULL ORDER BY created_at DESC").fetchall()
    conn.close()
    result = []
    for r in rows:
        job = dict(r)
        for k in ("stage1_outputs", "cascade_outputs", "assembled_files", "assembled_content"):
            try:
                job[k] = json.loads(job.get(k) or "{}")
            except Exception:
                job[k] = {}
        result.append(job)
    return result

def db_recover_missing_files():
    """Recreate output files from DB if they are missing (e.g. after redeploy)."""
    try:
        jobs = db_load_all_jobs()
    except Exception as e:
        logging.error("DB recovery failed: %s", e)
        return
    recovered = 0
    for job in jobs:
        for key, content in (job.get("assembled_content") or {}).items():
            fname = (job.get("assembled_files") or {}).get(key)
            if not fname or not content:
                continue
            path = OUTPUTS_DIR / fname
            if not path.exists():
                try:
                    path.write_text(content, encoding="utf-8")
                    recovered += 1
                except Exception as e:
                    logging.error("Recovery failed for %s: %s", fname, e)
    if recovered:
        logging.info("Recovered %d output files from DB", recovered)

init_db()
db_recover_missing_files()

app = FastAPI(title="Studio N")
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))

GITHUB_REPO         = os.getenv("GITHUB_REPO", "nickdcruz/nicklaus-marketing-agents")
GITHUB_TOKEN        = os.environ.get("GITHUB_TOKEN")
ANTHROPIC_API_KEY    = os.environ["ANTHROPIC_API_KEY"]
ARCADS_CLIENT_ID      = os.getenv("ARCADS_CLIENT_ID", os.getenv("ARCADS_API_KEY", ""))
ARCADS_CLIENT_SECRET  = os.getenv("ARCADS_CLIENT_SECRET", "")
ARCADS_BASE_URL       = os.getenv("ARCADS_BASE_URL", "https://external-api.arcads.ai")
ARCADS_CREDIT_BUDGET  = float(os.getenv("ARCADS_CREDIT_BUDGET", "200"))
HTTP_USER            = os.getenv("HTTP_USER", "admin")
HTTP_PASS           = os.getenv("HTTP_PASS", "changeme")
SESSION_SECRET      = os.getenv("SESSION_SECRET", secrets.token_hex(32))
PORT                = int(os.environ.get("PORT", "5050"))
APIFY_API_TOKEN     = os.environ.get("APIFY_API_TOKEN")
BUFFER_ACCESS_TOKEN = os.environ.get("BUFFER_ACCESS_TOKEN")
anthropic_client = AsyncAnthropic(api_key=ANTHROPIC_API_KEY)

# ── Arcads API helpers ────────────────────────────────────────────

def _arcads_headers():
    import base64
    creds = base64.b64encode(f"{ARCADS_CLIENT_ID}:{ARCADS_CLIENT_SECRET}".encode()).decode()
    return {"Authorization": f"Basic {creds}", "Content-Type": "application/json"}

async def arcads_get_products() -> list:
    if not ARCADS_CLIENT_ID:
        return []
    async with httpx.AsyncClient() as client:
        r = await client.get(f"{ARCADS_BASE_URL}/v1/products",
                             headers=_arcads_headers(), timeout=15)
        r.raise_for_status()
        data = r.json()
        if isinstance(data, list):
            return data
        return data.get("items") or data.get("data") or []

async def arcads_generate_video(payload: dict) -> dict:
    async with httpx.AsyncClient() as client:
        logging.info("Arcads generate payload: %s", json.dumps(payload))
        r = await client.post(f"{ARCADS_BASE_URL}/v2/videos/generate",
                              json=payload, headers=_arcads_headers(), timeout=60)
        logging.info("Arcads generate response %d: %s", r.status_code, r.text[:500])
        if not r.is_success:
            raise Exception(f"Arcads {r.status_code}: {r.text[:400]}")
        return r.json()

async def arcads_generate_image(payload: dict) -> dict:
    async with httpx.AsyncClient() as client:
        logging.info("Arcads image payload: %s", json.dumps(payload))
        r = await client.post(f"{ARCADS_BASE_URL}/v2/images/generate",
                              json=payload, headers=_arcads_headers(), timeout=30)
        logging.info("Arcads image response %d: %s", r.status_code, r.text[:400])
        if not r.is_success:
            raise Exception(f"Arcads {r.status_code}: {r.text[:300]}")
        return r.json()

async def arcads_poll_video(video_id: str) -> dict:
    """Poll job status via /v1/assets/{id} — used for all v2-generated jobs (video + image)."""
    async with httpx.AsyncClient() as client:
        r = await client.get(f"{ARCADS_BASE_URL}/v1/assets/{video_id}",
                             headers=_arcads_headers(), timeout=15)
        if not r.is_success:
            raise Exception(f"Asset poll {r.status_code}: {r.text[:200]}")
        return r.json()

async def arcads_poll_asset(asset_id: str) -> dict:
    async with httpx.AsyncClient() as client:
        r = await client.get(f"{ARCADS_BASE_URL}/v1/assets/{asset_id}",
                             headers=_arcads_headers(), timeout=15)
        if not r.is_success:
            raise Exception(f"Asset poll {r.status_code}: {r.text[:200]}")
        return r.json()

async def arcads_upload_file(file_bytes: bytes, filename: str, content_type: str) -> str:
    """Upload file to Arcads presigned storage, return filePath for use in generation requests."""
    async with httpx.AsyncClient() as client:
        # Step 1: get presigned URL
        r = await client.post(f"{ARCADS_BASE_URL}/v1/file-upload/get-presigned-url",
                              json={"fileType": content_type},
                              headers=_arcads_headers(), timeout=15)
        logging.info("Presigned URL response %d: %s", r.status_code, r.text[:300])
        if not r.is_success:
            raise Exception(f"Presigned URL failed {r.status_code}: {r.text[:200]}")
        data = r.json()
        upload_url = data.get("presignedUrl") or data.get("uploadUrl") or data.get("url", "")
        file_path  = data.get("filePath") or data.get("path", "")
        if not upload_url:
            raise Exception(f"No presigned URL in response: {data}")
        # Step 2: PUT file directly to S3 — no Arcads auth header, Content-Type must match fileType
        upload_r = await client.put(upload_url, content=file_bytes,
                                    headers={"Content-Type": content_type}, timeout=120)
        logging.info("S3 upload response %d (filePath=%s)", upload_r.status_code, file_path)
        if not upload_r.is_success:
            raise Exception(f"S3 upload failed {upload_r.status_code}: {upload_r.text[:200]}")
        return file_path

# SQLite table for video studio jobs
def init_video_db():
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS video_jobs (
            id          TEXT PRIMARY KEY,
            client      TEXT DEFAULT 'other',
            job_type    TEXT DEFAULT 'video',
            model       TEXT,
            prompt      TEXT,
            status      TEXT DEFAULT 'pending',
            arcads_id   TEXT,
            result_url  TEXT,
            formats     TEXT DEFAULT '{}',
            created_at  TEXT DEFAULT (datetime('now')),
            error       TEXT
        )
    """)
    conn.commit()
    conn.close()

init_video_db()

_agent_cache:       dict = {}
_jobs:              dict = {}
_completed:         dict = {}
_processing_status: dict = {}

VALID_AGENTS = {"callum", "priya", "dante", "suki", "felix", "nadia", "zara", "reeva", "kiara", "rex", "nova"}

# Seedance 2.0 takes "resolution" instead of "aspectRatio"; all ratios map to 720p.
SEEDANCE2_RESOLUTION_MAP: dict = {
    "16:9": "720p", "9:16": "720p",
    "1:1":  "720p", "4:3":  "720p",
}

# Per-model duration limits (seconds). None = no hard limit enforced by Arcads.
MODEL_DURATION_LIMITS: dict = {
    "kling-3.0":    (3, 15),
    "seedance-2.0": (3, 10),
    "sora2":        (1, 20),
    "veo31":        (1, 8),
    "grok-video":   (1, 30),
}

THIRD_PARTY_MESSAGES = {
    "video_production": (
        "Dante's video concept is saved and ready. Production requires Canva, CapCut, or Adobe Premiere — "
        "use real screen recordings of the software, not AI-generated imagery. Brief saved to outputs."
    ),
    "social_images": (
        "Suki's design brief is saved. Static images require real screen recordings or photography. "
        "Use the Canva JSON template in your outputs folder to build designs in Canva."
    ),
    "publishing": (
        "Content is approved. Connect Buffer, Hootsuite, or a platform API to enable auto-scheduling."
    ),
}

# ── Auth middleware ───────────────────────────────────────────────

PUBLIC_PATHS = {"/health", "/login"}

class SessionAuthMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        path = request.url.path
        if path in PUBLIC_PATHS or path.startswith("/api/stream/"):
            return await call_next(request)
        if not request.session.get("authenticated"):
            if path.startswith("/api/"):
                return JSONResponse({"error": "Not authenticated"}, status_code=401)
            return RedirectResponse("/login", status_code=302)
        return await call_next(request)

app.add_middleware(SessionAuthMiddleware)
app.add_middleware(SessionMiddleware, secret_key=SESSION_SECRET, session_cookie="studio_session",
                   max_age=86400 * 7, https_only=False, same_site="lax")

# ── Orchestration prompts ─────────────────────────────────────────

ORCHESTRATION_PREFIX = """SYSTEM INSTRUCTION — READ THIS BEFORE ANYTHING ELSE:
You are in orchestration mode. Your job is to analyse the brief and decide which specialist agents are needed.
After your analysis, end your response with a JSON block in this exact format — this is required:

```json
{"agents_needed": ["agent1", "agent2"], "briefs": {"agent1": "full brief text", "agent2": "full brief text"}}
```

VALID AGENT NAMES — USE ONLY THESE EXACT STRINGS. DO NOT INVENT OTHERS:
callum, priya, dante, suki, felix, nadia, zara, reeva, kiara, rex, nova

Agent specialisms:
- callum: LinkedIn and long-form written content
- priya: social media and short-form content (captions, hooks, threads)
- dante: video scripts, shot lists, reels direction — USE THIS for any video scripting task
- suki: static visual design briefs (Instagram posts, carousels, banners)
- felix: decks and presentations
- nadia: web copy and landing pages
- zara: research and competitive intelligence
- reeva: brand identity and strategy
- kiara: AI video generation via Arcads — USE THIS only when actually generating video with AI tools
- rex: Remotion programmatic video code
- nova: Nano Banana brand visuals and product stills

If a name is not on the list above, DO NOT use it. There is no agent called "lena", "luna", "leo", or any other name. Only the 11 names listed.

---
"""

REVIEW_CASCADE_SUFFIX = """

---
SYSTEM — REVIEW + CASCADE: After your verdicts, append ONE cascade block and nothing after it:

```cascade
{
  "verdicts": {"agent_name": "approved"},
  "cascade": {
    "next_agents": [],
    "assembly": "none",
    "flags": []
  }
}
```

ASSEMBLY GUIDE:
• Brand foundation delivered        → next_agents:[felix,nadia,suki]  assembly:html_website
• Website or landing page           → next_agents:[nadia,suki]        assembly:html_website
• One-pager or sales sheet          → next_agents:[felix,suki]        assembly:html_onepager
• Pitch deck or presentation        → next_agents:[felix]             assembly:html_deck
• Full social/content campaign      → next_agents:[priya,dante,suki]  assembly:social_pack  flags:[video_production,social_images]
• LinkedIn or long-form only        → next_agents:[]  assembly:none
• Research, strategy, or analysis   → next_agents:[]  assembly:none

flags options: video_production, social_images, publishing
"""

# ── Brand data extraction ─────────────────────────────────────────

def extract_brand_data(outputs: dict) -> dict:
    """Extract hex colors, font names, and brand name from any agent output."""
    combined = "\n".join(outputs.values())

    hex_colors = list(dict.fromkeys(re.findall(r"#[0-9A-Fa-f]{6}\b", combined)))[:6]

    google_fonts = [
        "Inter", "Plus Jakarta Sans", "DM Sans", "Sora", "Poppins",
        "Montserrat", "Raleway", "Lato", "Nunito", "Work Sans",
    ]
    found_fonts = [f for f in google_fonts if f.lower() in combined.lower()]
    chosen_fonts = found_fonts[:2] if found_fonts else ["Inter"]

    name_match = re.search(r"^#\s+(.+)$", combined, re.MULTILINE)
    brand_name = name_match.group(1).strip() if name_match else ""

    photo_match = re.search(
        r"(?:photography[:\s]+|photo\s+direction[:\s]+|visual\s+style[:\s]+|image\s+direction[:\s]+)"
        r"([^\n]{10,120})",
        combined, re.IGNORECASE)
    photography = photo_match.group(1).strip() if photo_match else \
        "Clean, professional product photography with natural lighting"

    return {
        "hex_colors":  hex_colors,
        "fonts":       chosen_fonts,
        "brand_name":  brand_name,
        "primary":     hex_colors[0] if hex_colors else "#1e293b",
        "secondary":   hex_colors[1] if len(hex_colors) > 1 else "#64748b",
        "accent":      hex_colors[2] if len(hex_colors) > 2 else "#3b82f6",
        "photography": photography,
    }

def _font_link(fonts: list[str]) -> str:
    families = "|".join(f.replace(" ", "+") + ":wght@300;400;600;700" for f in fonts)
    return f'<link href="https://fonts.googleapis.com/css2?family={families}&display=swap" rel="stylesheet">'

def _font_stack(fonts: list[str]) -> str:
    return ", ".join(f'"{f}"' for f in fonts) + ", -apple-system, sans-serif"

# ── Assembly prompts ──────────────────────────────────────────────

def _build_website_prompt(outputs: dict, brief: str) -> str:
    bd = extract_brand_data(outputs)
    p  = bd["primary"]
    s  = bd["secondary"]
    a  = bd["accent"]
    font_link  = _font_link(bd["fonts"])
    font_stack = _font_stack(bd["fonts"])
    inputs = "\n\n".join(f"## {k.upper()}\n\n{v}" for k, v in outputs.items() if v)

    return f"""You are a senior frontend engineer at a world-class design agency. Build a visually stunning, \
premium B2B marketing website as a single self-contained HTML file. \
Use only HTML, CSS, and inline SVG — zero external images, zero placeholder divs, \
zero dashed-border boxes. Every pixel must look intentionally designed.

BRAND PALETTE — use these exact hex values, no substitutions:
  Primary:   {p}
  Secondary: {s}
  Accent:    {a}
  Tint A:    {p}14   (primary at ~8% opacity — subtle section backgrounds)
  Tint B:    {p}28   (primary at ~16% opacity — cards, hover states)

TYPOGRAPHY:
  Include in <head>: {font_link}
  CSS font-family: {font_stack}

CONTENT INPUTS (all copy, features, and brand details come from here):
{inputs}

BRIEF: {brief}

━━━ SECTION-BY-SECTION VISUAL SPEC ━━━

Each section below has a mandatory visual treatment. Follow it exactly.

── NAV (id="top") ──
  Full-width bar. background: {p}; height: 68px.
  Max-width 1200px centred. Logo text: white, 700 weight, 17px.
  Nav links: color rgba(255,255,255,0.65); hover color white; font-size 14px.
  CTA button: background {a}; color white; border-radius 999px; padding 9px 22px; font-size 13px; font-weight 700; no border.
  Smooth underline animation on nav link hover using CSS ::after pseudo-element.

── HERO (id="hero") ──
  min-height: 100vh; position: relative; overflow: hidden.
  Background: radial-gradient(ellipse 130% 90% at 65% 50%, {p} 0%, {p}ee 45%, {s} 100%).
  Floating decorative shapes (position: absolute, pointer-events: none, z-index 0):
    — Large circle: width 700px, height 700px, border-radius 50%, background rgba(255,255,255,0.04), top: -15%, right: -10%
    — Medium circle: width 320px, height 320px, border-radius 50%, background {a}28, bottom: -8%, left: -5%
    — Rotated square: width 180px, height 180px, background rgba(255,255,255,0.03), transform rotate(45deg), top: 20%, right: 20%
  Two-column flex layout (align-items: center, gap: 60px, padding: 120px 8% 80px):
    LEFT COLUMN (flex 1.1):
      Eyebrow: text from inputs, uppercase, {a}, font-size 11px, letter-spacing 3px, font-weight 700, margin-bottom 18px.
      H1: main headline from copy inputs, white, font-size clamp(42px,5.5vw,72px), font-weight 800, line-height 1.08, margin-bottom 22px.
      Subtitle: white at 75% opacity, font-size 18px, line-height 1.65, max-width 480px, margin-bottom 36px.
      Two CTA buttons side-by-side (flex, gap 14px):
        Primary: background {a}, color white, padding 14px 28px, border-radius 10px, font-size 15px, font-weight 700.
        Secondary: background rgba(255,255,255,0.1), color white, border: 1px solid rgba(255,255,255,0.25), same padding.
    RIGHT COLUMN (flex 1): inline SVG product screen mockup, viewBox="0 0 520 360":
      Outer drop shadow: filter drop-shadow(0 32px 64px rgba(0,0,0,0.4)).
      Frame rect: x=0 y=0 w=520 h=360 rx=14, fill=#0f172a.
      Title bar: rect x=0 y=0 w=520 h=40 rx=14, fill={p}.
        Traffic-light circles: cx=20,32,44 cy=20 r=6, fills #ff5f57, #febc2e, #28c840.
        Title text: x=260 y=25, text-anchor=middle, fill=rgba(255,255,255,0.5), font-size=12.
      Content area (below y=40):
        — Sidebar: rect x=0 y=40 w=130 h=320, fill=rgba(0,0,0,0.2).
          3 sidebar nav items as rounded rects: y=60,96,132, x=12, w=108, h=24, rx=6;
          first one filled {a}40, others rgba(255,255,255,0.05). Small square icons (8×8) at x=20.
        — Main content (x>130):
          Metric cards row (y=56): 3 rects, each 90×56, rx=8, fill=rgba(255,255,255,0.06), x=148,256,364.
            Inside each: big number text in {a} (font-size=18, y=80), small label in rgba(255,255,255,0.4) (font-size=9, y=96).
          Bar chart (y=128 to 260): 6 vertical bars at x=152,178,204,230,256,282, width=18, rx=3.
            Heights vary: 80,110,60,130,90,120 (bars grow downward from y=260, so y = 260-height).
            Bars filled {a}; last bar slightly lighter ({a}99).
          Table rows (x=310, y=136): 4 rows, height=28 each, width=200, rx=4;
            alternating fill rgba(255,255,255,0.04) and transparent.
            Short text lines: rect w=60 h=6 rx=3 fill=rgba(255,255,255,0.2) and w=40 h=6 rx=3 fill={a}60.

── FEATURES (id="features") — value proposition ──
  background: white; padding: 100px 8%.
  Section header centred (margin-bottom 56px):
    Overline: "WHY CHOOSE US" (or relevant), {a}, 11px, letter-spacing 3px, font-weight 700.
    H2: brand proposition headline from inputs, {p}, font-size clamp(28px,3.5vw,42px), font-weight 800.
    Subtitle: 16px, color #64748b, max-width 560px, centred.
  3-column CSS grid (grid-template-columns: repeat(3,1fr); gap: 28px). Each card:
    background: white; border-radius: 20px; padding: 36px 28px;
    box-shadow: 0 2px 12px rgba(0,0,0,0.06), 0 0 0 1px rgba(0,0,0,0.04);
    border-top: 3px solid {a};
    transition: transform 0.2s, box-shadow 0.2s;
    On hover: transform translateY(-6px); box-shadow: 0 16px 48px rgba(0,0,0,0.12).
    Icon: inline SVG 52×52px — a circle fill={p}18, inside it a relevant geometric path/shape in {p} (e.g. checkmark, graph bars, gear spokes, lightning bolt — use actual SVG path data, not emoji).
    H3: {p}, font-size 17px, font-weight 700, margin: 18px 0 10px.
    Body: #64748b, font-size 14px, line-height 1.7.

── HOW IT WORKS (id="how-it-works") — features in depth ──
  background: {p}08; padding: 100px 8%.
  Section header centred (same style as above).
  2–3 alternating rows (flex, align-items centre, gap 72px, margin-bottom 80px).
  Odd rows: text left, visual right. Even rows: visual left, text right.
  TEXT SIDE (flex 1):
    Step badge: inline block, {a}, font-size 48px, font-weight 900, opacity 0.12, position absolute, top -10px, left -6px. (The number sits as a giant ghost behind the heading.)
    H3: {p}, 22px, 700.
    Body: #475569, 15px, line-height 1.7, margin-bottom 16px.
    Bullet list: each item has a 4px wide left border in {a} and padding-left 14px.
  VISUAL SIDE (flex 1): inline SVG "feature screen", viewBox="0 0 440 300":
    Container rect: rx=12, fill=#1e293b, w=440, h=300.
    Top bar: h=36, fill={p}, rx=12 (top corners).
      Three circles traffic-light style. Tab labels: 2–3 small rounded rects fill=rgba(255,255,255,0.15).
    Content: design a unique abstract UI for each feature step using rects, circles, polylines.
      Use {a} for highlighted elements, rgba(255,255,255,0.08) for row/card backgrounds.
      Make each SVG visually different — vary the layout between a table, a chart, a form, a map grid, etc.

── TESTIMONIALS (id="testimonials") ──
  background: {p}; padding: 90px 8%.
  Section header centred in white.
  2–3 testimonial cards in flex (gap 24px). Each card:
    background: rgba(255,255,255,0.07); border-radius: 20px; padding: 36px 32px; flex: 1;
    border: 1px solid rgba(255,255,255,0.1);
    Quote mark SVG (position: absolute, top: 20px, left: 24px): a large " mark, fill={a}, opacity 0.35, font-size 80px (use SVG text element).
    Quote text: white at 90% opacity, 15px, italic, line-height 1.7, padding-top 32px.
    Separator: 1px solid rgba(255,255,255,0.15), margin 20px 0.
    Avatar row: inline SVG circle (r=22, fill={a}33, stroke={a}, stroke-width=1.5) with initials text, plus name/title.
    Name: white, 13px, 700. Title: rgba(255,255,255,0.55), 12px.

── CTA (id="contact") ──
  position: relative; overflow: hidden; padding: 110px 8%; text-align: centre.
  Background: linear-gradient(135deg, {a} 0%, {p} 100%).
  Decorative rings (position: absolute, pointer-events none):
    — Circle 600px, border: 2px solid rgba(255,255,255,0.08), border-radius 50%, centred top-left -200px,-200px.
    — Circle 400px, same style, bottom-right -150px,-150px.
  H2: white, clamp(32px,4vw,52px), 800, margin-bottom 18px.
  Subtitle: rgba(255,255,255,0.8), 17px, margin-bottom 40px.
  Two buttons: Primary white bg + {p} text; Secondary transparent + white border + white text.

── FOOTER ──
  background: #0f172a; padding: 60px 8% 32px; color: rgba(255,255,255,0.45).
  3-column grid: brand+tagline+social icons | nav links | contact info.
  Brand name white 700. Links hover white. Divider line then copyright bar, font-size 12px.

━━━ TECHNICAL REQUIREMENTS ━━━

MANDATORY IDs — these are non-negotiable. Every section must carry exactly these id attributes.
Do NOT invent custom ids based on brand names or content. Use exactly these strings:
  <nav>  or  <header>        →  id="top"
  Hero section               →  id="hero"
  Value proposition section  →  id="features"
  How-it-works section       →  id="how-it-works"
  Testimonials section       →  id="testimonials"
  CTA / contact section      →  id="contact"

MANDATORY NAV HREFS — every nav link must use exactly these href values:
  href="#hero"  href="#features"  href="#how-it-works"  href="#testimonials"  href="#contact"

MANDATORY CSS — include in the <style> block:
  html {{ scroll-behavior: smooth; }}

- Single self-contained HTML5 file. Google Fonts CDN <link> in <head>. All CSS in one <style> block.
- Responsive: columns stack below 768px; nav links collapse on mobile.
- All SVG elements must be self-contained inline SVG with explicit viewBox. No <image> tags inside SVG.
- No external images. No placeholder divs. No dashed-border boxes. No emoji as icons.
- All content (headlines, copy, feature names, testimonials) must come from the marketing inputs above.
- You MUST generate ALL six sections plus footer before closing </html>. Do not stop early.

Output ONLY valid HTML. Start with <!DOCTYPE html>. End with </html>. Zero preamble, zero explanation."""

def _build_onepager_prompt(outputs: dict, brief: str) -> str:
    bd = extract_brand_data(outputs)
    p  = bd["primary"]
    s  = bd["secondary"]
    a  = bd["accent"]
    font_link  = _font_link(bd["fonts"])
    font_stack = _font_stack(bd["fonts"])
    inputs = "\n\n".join(f"## {k.upper()}\n\n{v}" for k, v in outputs.items() if v)
    return f"""You are a senior designer producing a premium B2B sales one-pager as a print-ready HTML file. \
This document goes in front of enterprise buyers and investors. \
It must look like a top-agency production piece. \
Use only HTML, CSS, and inline SVG — zero external images, zero placeholder divs, zero dashed boxes.

BRAND PALETTE:
  Primary:   {p}
  Secondary: {s}
  Accent:    {a}

TYPOGRAPHY:
  Include in <head>: {font_link}
  CSS font-family: {font_stack}

CONTENT INPUTS:
{inputs}

BRIEF: {brief}

━━━ PRINT LAYOUT SPEC — ONE A4 PAGE ━━━

The entire document must fit on a single A4 sheet (210mm × 297mm) when printed.
Use a fixed-width wrapper: width: 794px; margin: 0 auto; (794px ≈ A4 at 96 dpi).
No section should overflow. Tight spacing throughout.

── PRINT TOOLBAR (.no-print) ──
  Visible on screen, hidden on print. White bar above the page.
  "Save as PDF" button: background {p}; color white; border-radius 7px; padding 8px 20px; font-size 13px; font-weight 700.
  "Close" button: border 1px solid #e2e8f0; color #475569; same padding.

── HEADER BAND (full 794px width, height ~80px) ──
  background: linear-gradient(90deg, {p} 0%, {s} 100%); padding: 0 36px.
  Brand name: white, 20px, 800 weight. Tagline beside it: rgba(255,255,255,0.65), 12px.
  Right side: inline SVG logomark — a 40×40 geometric shape using {a} and white (abstract mark, not text).

── HERO STRIP (background {p}14, padding 28px 36px) ──
  Two columns (flex, align-items centre):
    Left (~58%): Value proposition headline from inputs, {p}, 26px, 800, line-height 1.2, max 2 lines.
    Right (~42%): One sentence supporting statement, #475569, 13px, line-height 1.6.

── STATS BAR (background {p}, padding 18px 36px) ──
  Flex row, justify-content space-evenly.
  Extract 3 concrete metrics from inputs (e.g. speed improvement, ROI, time-to-value).
  Each stat block: centred, divider lines between.
    Number: {a}, 30px, 800 weight.
    Label: rgba(255,255,255,0.65), 10px, uppercase, letter-spacing 1.5px.
  Thin vertical dividers: 1px solid rgba(255,255,255,0.18).

── MAIN BODY (padding 28px 36px, flex row, gap 28px) ──
  LEFT COLUMN (~52%, flex-direction column):
    "Why [BrandName]" heading: {p}, 14px, 700, letter-spacing 0.5px, margin-bottom 16px.
    3 feature blocks stacked (margin-bottom 18px each):
      Inline SVG icon (28×28): circle fill={p}18, geometric path inside in {p}.
      Feature title: {p}, 13px, 700, margin-bottom 5px.
      Feature body: #64748b, 12px, line-height 1.6.
      Left accent: border-left 3px solid {a}; padding-left 12px; on the title+body.
    Testimonial quote block (margin-top 20px):
      background {p}0d; border-radius 10px; padding 16px 18px;
      Large inline SVG open-quote mark, fill={a}, opacity 0.4, float left, width 20px.
      Quote text: #334155, 12px, italic, line-height 1.6.
      Attribution: {p}, 11px, 700, margin-top 8px.
  RIGHT COLUMN (~48%):
    Inline SVG product dashboard illustration, viewBox="0 0 340 280":
      Frame: rect w=340 h=280 rx=10 fill=#0f172a.
      Top bar: h=32 fill={p} rx=10 (top corners only — bottom is square).
        Traffic circles: cx=14,26,38 cy=16 r=5 fills #ff5f57, #febc2e, #28c840.
        Tab strip: 3 small rounded rects fill=rgba(255,255,255,0.1), w=48, h=16, y=8, x=60/116/172.
      Left sidebar: rect x=0 y=32 w=80 h=248 fill=rgba(0,0,0,0.18).
        4 nav items: rounded rects y=48,76,104,132 x=8 w=64 h=20 rx=5;
        first fill={a}50, rest rgba(255,255,255,0.05).
      Main panel (x>80):
        KPI cards row (y=42): 3 cards, each 70×44 rx=6 fill=rgba(255,255,255,0.06) x=90,172,254.
          Big number text {a} font-size=14 y=62; small label rgba(255,255,255,0.35) font-size=7 y=76.
        Chart area (y=96 to 200 x=88 to 330):
          Background rect fill=rgba(255,255,255,0.03) rx=6.
          5 vertical bars x=100,130,160,190,220 width=18 rx=3 fill={a}; heights 60,85,45,100,70 (from y=200 upward).
          Line overlay: polyline points connecting bar tops, stroke={a}80 stroke-width=1.5 fill=none.
        Table rows (y=212 to 272): 3 rows h=18 w=240 x=88 rx=3;
          alternating fill rgba(255,255,255,0.03)/transparent.
          Two text stubs per row: w=50 h=5 rx=2 and w=30 h=5 rx=2, fill=rgba(255,255,255,0.15) and {a}60.

── FOOTER BAR (background {p}, height 36px, padding 0 36px) ──
  Flex, align-items centre, justify-content space-between.
  Brand name: white, 11px, 700. Centre: website URL placeholder: rgba(255,255,255,0.6), 10px. Right: "Confidential": rgba(255,255,255,0.4), 10px.

━━━ TECHNICAL REQUIREMENTS ━━━
- Self-contained HTML5. Google Fonts CDN in <head>. All CSS in one <style> block.
- @page {{ size: A4; margin: 0; }}
- body {{ margin: 0; background: #f1f5f9; }}
- .page {{ width: 794px; margin: 20px auto; background: white; }}
- @media print {{ .no-print {{ display: none !important; }} body {{ background: white; margin: 0; }} .page {{ margin: 0; box-shadow: none; }} }}
- Zero external images. Zero placeholder divs. Zero dashed boxes. Every zone fully designed.
- All content (metrics, features, testimonial, brand name) must come from the inputs above.

Output ONLY valid HTML. Start with <!DOCTYPE html>. No explanation."""

def _build_deck_prompt(outputs: dict, brief: str) -> str:
    bd = extract_brand_data(outputs)
    font_link = _font_link(bd["fonts"])
    font_stack = _font_stack(bd["fonts"])
    inputs = "\n\n".join(f"## {k.upper()}\n\n{v}" for k, v in outputs.items() if v)
    return f"""Generate a complete HTML presentation deck for a B2B product pitch.

BRAND COLORS: Primary: {bd['primary']} | Secondary: {bd['secondary']} | Accent: {bd['accent']}
FONTS CDN: {font_link}
Font stack: {font_stack}

INPUTS:
{inputs}

ORIGINAL BRIEF: {brief}

REQUIREMENTS:
- Self-contained HTML5, Google Fonts CDN in <head>, all CSS inline.
- Each slide is a full-viewport <section class="slide"> div.
- 10–12 slides: Title, Problem, Solution, Key Features (3), Demo slide (demo-placeholder: "📹 Live demo / screen recording"), Proof/Results, Pricing/Plans, Team placeholder, CTA/Next Steps.
- Keyboard navigation (ArrowLeft/ArrowRight) + on-screen prev/next buttons.
- Slide counter (e.g. "3 / 12") in corner.
- Brand color scheme: header/title slides use primary color background; content slides white with primary accents.
- demo-placeholder on demo slide: background:rgba(255,255,255,0.15); border:2px dashed rgba(255,255,255,0.5); border-radius:10px; padding:60px; text-align:center; color:rgba(255,255,255,0.8); font-size:16px;
- @media print: each slide prints as one page.

Output ONLY valid HTML. Start with <!DOCTYPE html>. No explanation."""

async def assemble_html(prompt: str, system_msg: str) -> str:
    html = await call_agent(system_msg, prompt, model="claude-opus-4-6", max_tokens=16000)
    m = re.search(r"<!DOCTYPE", html, re.IGNORECASE)
    return html[m.start():] if m else html

async def assemble_social_pack(outputs: dict, brief: str) -> str:
    inputs = "\n\n".join(f"## {k.upper()}\n\n{v}" for k, v in outputs.items() if v)
    prompt = f"""Compile these marketing outputs into a structured Social Media Content Pack.

INPUTS:\n{inputs}\n\nORIGINAL BRIEF: {brief}

Format: clean markdown with sections: Cover (client, date, brief summary) | LinkedIn Posts (numbered, copy-paste ready) | Instagram Captions (numbered, with hashtag sets) | TikTok/Reels Scripts | Facebook Posts | Design Briefs for Static Posts | Video Production Brief (flagged: use real screen recordings, not AI images) | Hashtag Master List.

Output clean final markdown. No meta-commentary."""
    return await call_agent("You compile marketing content packs.", prompt,
                            model="claude-opus-4-6", max_tokens=4096)

ASSEMBLERS = {
    "html_website":        (".html",  "website"),
    "html_onepager":       (".html",  "one-pager"),
    "social_pack":         (".md",    "social-pack"),
    "canva_json":          (".json",  "canva"),
    "social_pack_cascade": (".md",    "social-pack"),
    "video_brief":         (".html",  "video-brief"),
}

# ── Canva JSON template ───────────────────────────────────────────

def _build_canva_json_prompt(outputs: dict, brief: str) -> str:
    bd = extract_brand_data(outputs)
    inputs = "\n\n".join(f"## {k.upper()}\n\n{v}" for k, v in outputs.items() if v)
    palette_json = json.dumps(bd["hex_colors"])
    hfont = bd["fonts"][0]
    bfont = bd["fonts"][-1] if len(bd["fonts"]) > 1 else bd["fonts"][0]
    return f"""Extract all approved brand and copy data into a structured JSON template for Canva.

BRAND COLORS:
Primary: {bd['primary']}
Secondary: {bd['secondary']}
Accent: {bd['accent']}
Full palette: {', '.join(bd['hex_colors'])}

FONTS: {', '.join(bd['fonts'])}

APPROVED MARKETING OUTPUTS:
{inputs}

ORIGINAL BRIEF: {brief}

Generate a JSON object with this exact structure — fill ALL fields with real extracted content:

{{
  "brand": {{
    "name": "extracted brand name",
    "tagline": "extracted tagline",
    "description": "one paragraph brand description",
    "url": "website URL if mentioned, else empty string",
    "colors": {{
      "primary": "{bd['primary']}",
      "secondary": "{bd['secondary']}",
      "accent": "{bd['accent']}",
      "palette": {palette_json}
    }},
    "typography": {{
      "heading_font": "{hfont}",
      "body_font": "{bfont}"
    }}
  }},
  "copy": {{
    "headline": "main headline from approved copy",
    "subheadline": "subheadline",
    "value_proposition": "one sentence value prop",
    "cta_primary": "primary CTA button text",
    "cta_secondary": "secondary CTA text",
    "features": [
      {{"title": "feature 1 name", "description": "brief description"}},
      {{"title": "feature 2 name", "description": "brief description"}},
      {{"title": "feature 3 name", "description": "brief description"}}
    ],
    "testimonials": [
      {{"quote": "testimonial text", "author": "Name, Title, Company"}},
      {{"quote": "second testimonial", "author": "Name, Title, Company"}}
    ],
    "social_posts": [
      {{"platform": "linkedin", "caption": "copy-paste ready 150-word post"}},
      {{"platform": "instagram", "caption": "short caption with 10-15 hashtags"}},
      {{"platform": "facebook", "caption": "post copy"}}
    ]
  }},
  "canva_bulk_create": {{
    "note": "Import these rows into Canva Bulk Create. Map field names to named elements in your template.",
    "rows": [
      {{"slide": "Hero", "Headline": "main headline", "Subheadline": "subheadline", "Body": "supporting line", "CTA": "button text", "Color1": "{bd['primary']}", "Color2": "{bd['secondary']}", "Font": "{hfont}"}},
      {{"slide": "Value Prop", "Headline": "section header", "Body": "value proposition paragraph", "Color1": "{bd['primary']}", "Color2": "{bd['secondary']}", "Font": "{hfont}"}},
      {{"slide": "Feature 1", "Headline": "feature 1 title", "Body": "feature 1 description", "Color1": "{bd['primary']}", "Color2": "{bd['secondary']}", "Font": "{hfont}"}},
      {{"slide": "Feature 2", "Headline": "feature 2 title", "Body": "feature 2 description", "Color1": "{bd['primary']}", "Color2": "{bd['secondary']}", "Font": "{hfont}"}},
      {{"slide": "Feature 3", "Headline": "feature 3 title", "Body": "feature 3 description", "Color1": "{bd['primary']}", "Color2": "{bd['secondary']}", "Font": "{hfont}"}},
      {{"slide": "Testimonial", "Headline": "quote text", "Body": "author attribution", "Color1": "{bd['primary']}", "Color2": "{bd['secondary']}", "Font": "{hfont}"}},
      {{"slide": "CTA", "Headline": "closing headline", "Body": "supporting close", "CTA": "final button text", "Color1": "{bd['primary']}", "Color2": "{bd['secondary']}", "Font": "{hfont}"}}
    ]
  }}
}}

Output ONLY valid JSON. Start with {{ and end with }}. No markdown, no explanation."""

async def assemble_canva_json(outputs: dict, brief: str) -> str:
    prompt = _build_canva_json_prompt(outputs, brief)
    raw = await call_agent(
        "You extract brand and copy data into structured JSON for Canva. Output only valid JSON.",
        prompt, model="claude-sonnet-4-6", max_tokens=4096)
    m = re.search(r"\{[\s\S]*\}", raw)
    if m:
        try:
            return json.dumps(json.loads(m.group(0)), indent=2, ensure_ascii=False)
        except json.JSONDecodeError:
            pass
    return raw

# ── Social cascade + video brief assembly ─────────────────────────

async def assemble_social_cascade(outputs: dict, brief: str) -> str:
    """Fire Callum (LinkedIn) + Priya (Instagram/Facebook) in parallel, return social pack markdown."""
    all_content = "\n\n".join(f"## {k.upper()}\n\n{v}" for k, v in outputs.items() if v)
    social_brief = (
        f"Original brief: {brief}\n\n"
        "APPROVED MARKETING OUTPUTS (use as source of truth for all content):\n\n"
        f"{all_content}\n\n---\n\n"
        "Produce one piece of final, copy-paste-ready social content. "
        "Make it specific to the brand and campaign above — no filler, no placeholders."
    )

    async def run(name: str) -> tuple[str, str]:
        sys_ = await fetch_agent(name)
        out  = await call_agent(sys_, social_brief)
        return name, out

    results = await asyncio.gather(run("callum"), run("priya"))
    agent_outputs = dict(results)
    callum_out = agent_outputs.get("callum", "")
    priya_out  = agent_outputs.get("priya", "")

    ts = datetime.now().strftime("%Y-%m-%d %H:%M")
    return (
        f"# Social Content Pack\n\n"
        f"**Generated:** {ts}  \n"
        f"**Brief:** {brief}\n\n"
        "---\n\n"
        "## LinkedIn Post (Callum)\n\n"
        f"{callum_out}\n\n"
        "---\n\n"
        "## Instagram & Facebook (Priya)\n\n"
        f"{priya_out}\n"
    )


async def assemble_video_brief(outputs: dict, brief: str) -> str:
    """Fire Dante for a video brief, return PDF-ready HTML."""
    all_content = "\n\n".join(f"## {k.upper()}\n\n{v}" for k, v in outputs.items() if v)
    bd = extract_brand_data(outputs)
    video_brief_brief = (
        f"Original brief: {brief}\n\n"
        "APPROVED MARKETING OUTPUTS:\n\n"
        f"{all_content}\n\n---\n\n"
        "Produce a complete, formatted video brief document. Include:\n"
        "1. Campaign overview and creative concept\n"
        "2. Shot list (numbered, each with shot type, location/setting, action, duration)\n"
        "3. Script outline (hook → problem → solution → CTA structure)\n"
        "4. TikTok specs (aspect ratio 9:16, duration 15–60s, caption, hashtags, hook text)\n"
        "5. Instagram Reels specs (aspect ratio 9:16, duration 15–90s, cover frame, caption)\n"
        "6. Hook direction (first 3 seconds — what grabs attention, exact opening line)\n"
        "7. Talent and production notes\n\n"
        "Use the brand colors and typography from the approved outputs above. "
        "Be specific — this brief goes straight to the production team.\n\n"
        "CRITICAL: You MUST deliver complete sections 1–7 for EVERY reel requested in the brief. "
        "Do not truncate, summarise, or skip any reel. If the brief asks for 5 reels, deliver all 5 in full."
    )
    sys_ = await fetch_agent("dante")
    content = await call_agent(sys_, video_brief_brief, model="claude-sonnet-4-6", max_tokens=16000)

    # Wrap in print-ready HTML
    primary   = bd["primary"]
    secondary = bd["secondary"]
    accent    = bd["accent"]
    font_link = _font_link(bd["fonts"])
    font_stack = _font_stack(bd["fonts"])
    body_html = _md_to_html(content)
    ts = datetime.now().strftime("%Y-%m-%d %H:%M")

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>Video Brief — {bd.get('brand_name', 'Campaign')}</title>
  {font_link}
  <style>
    * {{ box-sizing: border-box; margin: 0; padding: 0; }}
    body {{ font-family: {font_stack}; background: #f8fafc; color: #1e293b; padding: 0; }}
    .toolbar {{ background: {primary}; padding: 12px 32px; display: flex; gap: 10px; align-items: center; }}
    .toolbar-brand {{ color: white; font-size: 14px; font-weight: 700; flex: 1; }}
    .btn {{ background: white; color: {primary}; border: none; padding: 8px 18px; border-radius: 6px;
            font-size: 12px; font-weight: 700; cursor: pointer; }}
    .btn-outline {{ background: transparent; color: white; border: 1px solid rgba(255,255,255,0.4); }}
    .doc {{ max-width: 820px; margin: 0 auto; padding: 48px 40px 80px; background: white;
            min-height: 100vh; box-shadow: 0 0 40px rgba(0,0,0,0.06); }}
    .doc-header {{ border-left: 4px solid {primary}; padding-left: 20px; margin-bottom: 36px; }}
    .doc-header h1 {{ font-size: 22px; font-weight: 800; color: {primary}; margin-bottom: 4px; }}
    .doc-meta {{ font-size: 12px; color: #64748b; }}
    .brief-box {{ background: {primary}0d; border: 1px solid {primary}30; border-radius: 8px;
                  padding: 12px 16px; margin-bottom: 32px; font-size: 13px; color: #475569; }}
    h1 {{ font-size: 20px; font-weight: 800; color: {primary}; margin: 32px 0 12px; }}
    h2 {{ font-size: 17px; font-weight: 700; color: {primary}; margin: 28px 0 10px;
          padding-bottom: 6px; border-bottom: 1px solid {primary}20; }}
    h3 {{ font-size: 14px; font-weight: 700; color: #334155; margin: 18px 0 8px; }}
    p  {{ font-size: 14px; line-height: 1.8; margin-bottom: 12px; color: #334155; }}
    ul {{ padding-left: 20px; margin-bottom: 14px; }}
    li {{ font-size: 14px; line-height: 1.7; margin-bottom: 5px; color: #334155; }}
    strong {{ font-weight: 700; color: #1e293b; }}
    hr {{ border: none; border-top: 1px solid #e2e8f0; margin: 24px 0; }}
    @media print {{
      .toolbar {{ display: none !important; }}
      body {{ background: white; }}
      .doc {{ box-shadow: none; padding: 20px; }}
    }}
  </style>
</head>
<body>
  <div class="toolbar no-print">
    <div class="toolbar-brand">📹 Video Brief</div>
    <button class="btn" onclick="window.print()">Save as PDF</button>
    <button class="btn btn-outline" onclick="window.close()">Close</button>
  </div>
  <div class="doc">
    <div class="doc-header">
      <h1>Video Production Brief</h1>
      <div class="doc-meta">Generated {ts}</div>
    </div>
    <div class="brief-box"><strong>Brief:</strong> {brief}</div>
    {body_html}
  </div>
</body>
</html>"""


# ── Apify / Zara competitor enrichment ───────────────────────────

_COMPETITOR_RE = re.compile(
    r"\b(competitor[s]?|competitive|intelligence|benchmark|audit|landscape|market research|rival[s]?)\b",
    re.IGNORECASE,
)

def _extract_handles(text: str) -> list[str]:
    """Pull @handles and names after 'competitor/rival/vs' from text."""
    handles = re.findall(r"@([\w.]{2,30})", text)
    names   = re.findall(
        r"(?:competitor[s]?|rival[s]?|vs\.?)[:\s]+([A-Za-z0-9_.]{2,30})",
        text, re.IGNORECASE,
    )
    seen: dict[str, None] = {}
    for h in handles + names:
        seen[h.lstrip("@").lower()] = None
    return list(seen.keys())

def _run_apify_sync(handles: list[str]) -> dict:
    client = ApifyClient(APIFY_API_TOKEN)
    run = client.actor("apify/instagram-scraper").call(run_input={
        "directUrls": [f"https://www.instagram.com/{h}/" for h in handles],
        "resultsType": "posts",
        "resultsLimit": 5,
    })
    items = list(client.dataset(run["defaultDatasetId"]).iterate_items())
    return {"handles": handles, "posts": items[:20]}

async def enrich_zara_with_apify(zara_output: str, brief: str) -> str:
    """Append live Instagram competitor data to Zara's output when brief signals research."""
    if not APIFY_API_TOKEN:
        return zara_output
    if not _COMPETITOR_RE.search(brief):
        return zara_output
    handles = _extract_handles(brief + "\n" + zara_output)
    if not handles:
        return zara_output
    logging.info("Apify: scraping Instagram for handles %s", handles)
    try:
        loop = asyncio.get_event_loop()
        data = await loop.run_in_executor(None, _run_apify_sync, handles)
    except Exception as e:
        logging.error("Apify scrape failed: %s", e)
        return zara_output
    posts = data.get("posts", [])
    if not posts:
        return zara_output
    lines = ["\n\n---\n\n## Competitor Intelligence — Live Instagram Data\n"]
    for post in posts[:10]:
        handle  = post.get("ownerUsername", "unknown")
        caption = (post.get("caption") or "")[:200].replace("\n", " ")
        likes   = post.get("likesCount", 0)
        comments = post.get("commentsCount", 0)
        url     = post.get("url", "")
        lines.append(f"**@{handle}** — {likes} likes · {comments} comments")
        if caption:
            lines.append(f"> {caption.strip()}")
        if url:
            lines.append(f"[View post]({url})")
        lines.append("")
    return zara_output + "\n".join(lines)


# ── Buffer integration ────────────────────────────────────────────

def _post_to_buffer_sync(text: str) -> dict:
    profiles_r = requests.get(
        "https://api.bufferapp.com/1/profiles.json",
        params={"access_token": BUFFER_ACCESS_TOKEN},
        timeout=15,
    )
    profiles_r.raise_for_status()
    profiles = profiles_r.json()
    linkedin = [p for p in profiles if p.get("service") == "linkedin"]
    target   = linkedin if linkedin else profiles
    if not target:
        raise ValueError("No Buffer profiles found")
    profile_id = target[0]["id"]
    update_r = requests.post(
        "https://api.bufferapp.com/1/updates/create.json",
        data={
            "access_token": BUFFER_ACCESS_TOKEN,
            "profile_ids[]": profile_id,
            "text": text[:3000],
        },
        timeout=15,
    )
    update_r.raise_for_status()
    return update_r.json()

async def post_to_buffer(social_pack_md: str) -> None:
    if not BUFFER_ACCESS_TOKEN:
        return
    m = re.search(
        r"## LinkedIn Post \(Callum\)\s*\n\n(.+?)(?:\n\n---|\Z)",
        social_pack_md, re.DOTALL,
    )
    if not m:
        logging.warning("Buffer: LinkedIn post not found in social pack")
        return
    linkedin_text = m.group(1).strip()
    if not linkedin_text:
        return
    logging.info("Buffer: queuing LinkedIn post (%d chars)", len(linkedin_text))
    try:
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(None, _post_to_buffer_sync, linkedin_text)
        logging.info("Buffer: queued successfully — %s", result)
    except Exception as e:
        logging.error("Buffer post failed: %s", e)


# ── Campaign system ────────────────────────────────────────────────

_campaign_status: dict = {}  # campaign_id → {step, message, calendar?}

def init_campaign_db():
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS campaigns (
            id TEXT PRIMARY KEY,
            client TEXT DEFAULT 'other',
            platform TEXT DEFAULT 'instagram',
            duration_days INTEGER DEFAULT 30,
            style_brief TEXT DEFAULT '',
            reference_accounts TEXT DEFAULT '[]',
            status TEXT DEFAULT 'processing',
            current_step TEXT DEFAULT '',
            calendar_json TEXT DEFAULT '{}',
            video_job_ids TEXT DEFAULT '[]',
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    """)
    conn.commit()
    conn.close()

init_campaign_db()

# Multi-platform Apify scrapers

def _run_apify_instagram_sync(handles: list, max_posts: int = 25) -> list:
    client = ApifyClient(APIFY_API_TOKEN)
    run = client.actor("apify/instagram-scraper").call(run_input={
        "directUrls": [f"https://www.instagram.com/{h.lstrip('@')}/" for h in handles],
        "resultsType": "posts",
        "resultsLimit": max_posts,
    })
    return list(client.dataset(run["defaultDatasetId"]).iterate_items())[:max_posts]

def _run_apify_tiktok_sync(handles: list, max_posts: int = 25) -> list:
    client = ApifyClient(APIFY_API_TOKEN)
    run = client.actor("clockworks/free-tiktok-scraper").call(run_input={
        "profiles": [h.lstrip("@") for h in handles],
        "resultsPerPage": max_posts,
    })
    return list(client.dataset(run["defaultDatasetId"]).iterate_items())[:max_posts]

def _run_apify_youtube_sync(handles: list, max_videos: int = 15) -> list:
    client = ApifyClient(APIFY_API_TOKEN)
    urls = [{"url": f"https://www.youtube.com/@{h.lstrip('@')}"} for h in handles]
    run = client.actor("streamers/youtube-scraper").call(run_input={
        "startUrls": urls,
        "maxResults": max_videos,
    })
    return list(client.dataset(run["defaultDatasetId"]).iterate_items())[:max_videos]

def _run_apify_facebook_sync(handles: list, max_posts: int = 20) -> list:
    client = ApifyClient(APIFY_API_TOKEN)
    run = client.actor("apify/facebook-pages-scraper").call(run_input={
        "startUrls": [{"url": f"https://www.facebook.com/{h.lstrip('@')}"} for h in handles],
        "maxPosts": max_posts,
    })
    return list(client.dataset(run["defaultDatasetId"]).iterate_items())[:max_posts]

async def run_platform_research(platform: str, handles: list, urls: list = None) -> dict:
    """Scrape the given platform for competitor/reference account posts."""
    if not APIFY_API_TOKEN or not handles:
        return {"platform": platform, "posts": [], "skipped": True}
    loop = asyncio.get_event_loop()
    try:
        if platform == "instagram":
            posts = await loop.run_in_executor(None, _run_apify_instagram_sync, handles, 25)
        elif platform == "tiktok":
            posts = await loop.run_in_executor(None, _run_apify_tiktok_sync, handles, 25)
        elif platform == "youtube":
            posts = await loop.run_in_executor(None, _run_apify_youtube_sync, handles, 15)
        elif platform == "facebook":
            posts = await loop.run_in_executor(None, _run_apify_facebook_sync, handles, 20)
        else:
            posts = await loop.run_in_executor(None, _run_apify_instagram_sync, handles, 25)
        return {"platform": platform, "posts": posts, "handles_scraped": handles}
    except Exception as e:
        logging.error("Apify %s research failed: %s", platform, e)
        return {"platform": platform, "posts": [], "error": str(e)}

def format_research_for_agents(research_data: dict) -> str:
    """Convert Apify post data into a readable string for agent context."""
    platform = research_data.get("platform", "unknown")
    posts    = research_data.get("posts", [])
    if research_data.get("skipped") or not posts:
        return f"No live scrape data for {platform}. Research from industry knowledge only."
    lines = [f"## Live {platform.capitalize()} Data — {len(posts)} recent posts scraped\n"]
    for p in posts[:20]:
        handle   = (p.get("ownerUsername") or p.get("authorMeta", {}).get("name", "") or
                    p.get("pageName", "unknown"))
        caption  = (p.get("caption") or p.get("text") or p.get("description") or
                    p.get("title") or "")[:200]
        likes    = (p.get("likesCount") or p.get("diggCount") or
                    p.get("statistics", {}).get("likeCount", 0) or 0)
        views    = (p.get("videoViewCount") or p.get("playCount") or
                    p.get("statistics", {}).get("viewCount", 0) or 0)
        url      = p.get("url") or p.get("webVideoUrl") or ""
        lines.append(f"**@{handle}** — {likes:,} likes · {views:,} views")
        if caption:
            lines.append(f"> {caption.strip()}")
        if url:
            lines.append(f"[View]({url})")
        lines.append("")
    return "\n".join(lines)

# Kiara → Arcads auto-fire bridge

def extract_video_specs_from_kiara(text: str) -> list:
    """Parse Kiara's output for ```json blocks that are Arcads job specs."""
    specs = []
    for m in re.finditer(r'```(?:json)?\s*(\{[\s\S]+?\})\s*```', text):
        try:
            spec = json.loads(m.group(1))
            if "model" in spec and "prompt" in spec:
                specs.append(spec)
        except Exception:
            pass
    return specs

async def fire_video_jobs(specs: list, client: str, campaign_id: str = "") -> list:
    """Submit Kiara specs to Arcads, persist in video_jobs table, return arcads IDs."""
    job_ids = []
    for spec in specs:
        try:
            model = spec.get("model", "kling-3.0")
            clean = dict(spec)
            # Normalise Seedance 2.0 — uses resolution not aspectRatio
            if model == "seedance-2.0" and "aspectRatio" in clean:
                clean["resolution"] = SEEDANCE2_RESOLUTION_MAP.get(clean.pop("aspectRatio"), "720p")
            result    = await arcads_generate_video(clean)
            arcads_id = result.get("id") or result.get("videoId", "")
            if not arcads_id:
                continue
            fmt = clean.get("aspectRatio") or clean.get("resolution") or "9:16"
            tag = f"[Campaign {campaign_id}] " if campaign_id else ""
            conn = sqlite3.connect(DB_PATH)
            conn.execute("""
                INSERT INTO video_jobs
                    (id, client, job_type, model, prompt, status, arcads_id, formats)
                VALUES (?,?,?,?,?,?,?,?)
            """, (
                str(uuid.uuid4()), client, "video", model,
                f"{tag}{clean.get('prompt','')[:400]}",
                "pending", arcads_id,
                json.dumps([{"format": fmt, "arcadsId": arcads_id}])
            ))
            conn.commit()
            conn.close()
            job_ids.append(arcads_id)
            logging.info("Campaign video job fired: model=%s arcads_id=%s", model, arcads_id)
        except Exception as e:
            logging.error("fire_video_jobs spec failed: %s", e)
    return job_ids

# Campaign pipeline

async def run_campaign(
    campaign_id: str, client: str, platform: str,
    duration_days: int, style_brief: str,
    reference_accounts: list, reference_urls: list, content_mix: dict,
):
    def set_step(step: str, msg: str = ""):
        _campaign_status[campaign_id] = {"step": step, "message": msg}
        try:
            conn = sqlite3.connect(DB_PATH)
            conn.execute("UPDATE campaigns SET current_step=? WHERE id=?", (f"{step}: {msg}"[:200], campaign_id))
            conn.commit()
            conn.close()
        except Exception:
            pass
        logging.info("Campaign %s — %s: %s", campaign_id[:8], step, msg)

    try:
        # ── Step 1: Research ──────────────────────────────────────
        set_step("research", f"Zara scraping {platform} competitor data...")
        research_data  = await run_platform_research(platform, reference_accounts, reference_urls)
        research_text  = format_research_for_agents(research_data)

        zara_sys   = await fetch_agent("zara")
        zara_brief = (
            f"CAMPAIGN RESEARCH REQUEST\n\nClient: {client} | Platform: {platform.upper()} | "
            f"Duration: {duration_days} days\nStyle: {style_brief}\n"
            f"Reference accounts: {', '.join(reference_accounts) or 'none specified'}\n\n"
            f"LIVE DATA SCRAPED:\n{research_text}\n\n"
            f"Produce a structured research report covering:\n"
            f"1. Top 5 viral content formats on {platform} right now for this client's industry\n"
            f"2. Best hook patterns — opening lines that stop the scroll\n"
            f"3. Optimal posting frequency and best times for {platform}\n"
            f"4. Content mix recommendation for {duration_days} days "
            f"(target ~{int(duration_days * content_mix.get('video', 0.3))} videos, "
            f"~{int(duration_days * content_mix.get('static', 0.5))} static)\n"
            f"5. Top hashtags and keywords\n6. Key competitor insight\n\n"
            f"Be specific and actionable. This feeds directly into the content calendar."
        )
        zara_analysis = await call_agent(zara_sys, zara_brief, max_tokens=4000)

        # ── Step 2: Calendar plan ─────────────────────────────────
        set_step("planning", "Marcus building the content calendar...")
        marcus_sys = await fetch_agent("marcus")

        n_video    = max(1, int(duration_days * content_mix.get("video", 0.3)))
        n_static   = max(1, int(duration_days * content_mix.get("static", 0.5)))
        n_carousel = max(0, duration_days - n_video - n_static)

        cal_brief = (
            f"CAMPAIGN CALENDAR REQUEST\n\nClient: {client} | Platform: {platform.upper()} | "
            f"{duration_days} days\nStyle: {style_brief}\n"
            f"Mix: {n_video} videos + {n_static} static posts + {n_carousel} carousels\n\n"
            f"ZARA RESEARCH:\n{zara_analysis}\n\n"
            f"Build a complete {duration_days}-day content calendar.\n"
            f"Output as a JSON array inside a ```json block — one object per day:\n"
            f'{{"day":1,"content_type":"video|static|carousel","theme":"...","hook":"...","key_message":"...","cta":"...","notes":"..."}}\n\n'
            f"ALL {duration_days} days must be in the array. No truncation.\n\n"
            f"Below the JSON, add a Campaign Strategy section with: overall narrative arc, "
            f"key brand messages, and platform-specific tips for {platform}."
        )
        cal_response  = await call_agent(marcus_sys, cal_brief, model="claude-opus-4-6", max_tokens=8000)

        calendar_days = []
        cal_m = re.search(r'```json\s*(\[[\s\S]+?\])\s*```', cal_response)
        if cal_m:
            try:
                calendar_days = json.loads(cal_m.group(1))
            except Exception:
                pass
        if not calendar_days:
            # Fall back — create skeleton
            for d in range(1, duration_days + 1):
                ct = "video" if d % 3 == 0 else ("carousel" if d % 5 == 0 else "static")
                calendar_days.append({"day": d, "content_type": ct, "theme": f"Day {d}", "hook": "", "key_message": "", "cta": "", "notes": ""})

        # ── Step 3: Captions + scripts in parallel ────────────────
        set_step("writing", "Priya writing captions, Dante scripting videos...")
        priya_sys = await fetch_agent("priya")
        dante_sys = await fetch_agent("dante")

        video_days = [d for d in calendar_days if d.get("content_type") == "video"]

        priya_msg = (
            f"Write ALL social captions for this {duration_days}-day {platform} campaign.\n"
            f"Client: {client} | Style: {style_brief}\n\n"
            f"CALENDAR:\n{json.dumps(calendar_days, indent=2)}\n\n"
            f"For every day, output a JSON object: "
            f'{{"day":N,"caption":"...","hashtags":["..."],"cta":"..."}}\n'
            f"Return a JSON array of all {duration_days} objects inside a ```json block."
        )
        dante_msg = (
            f"Write video scripts for ALL {len(video_days)} video days in this campaign.\n"
            f"Client: {client} | Platform: {platform} | Style: {style_brief}\n\n"
            f"VIDEO DAYS:\n{json.dumps(video_days, indent=2)}\n\n"
            f"For each video, output: "
            f'{{"day":N,"hook":"5-8 word opening","script":"15-30 sec script","on_screen_text":"...","visual_direction":"..."}}\n'
            f"Return a JSON array inside a ```json block."
        )

        priya_out, dante_out = await asyncio.gather(
            call_agent(priya_sys, priya_msg, max_tokens=6000),
            call_agent(dante_sys, dante_msg, max_tokens=6000),
        )

        captions, scripts = [], []
        for raw, store in [(priya_out, captions), (dante_out, scripts)]:
            m = re.search(r'```json\s*(\[[\s\S]+?\])\s*```', raw)
            if m:
                try:
                    store.extend(json.loads(m.group(1)))
                except Exception:
                    pass

        # ── Step 4: Kiara video specs ─────────────────────────────
        set_step("video_specs", f"Kiara building Arcads specs for {len(video_days)} videos...")
        kiara_sys = await fetch_agent("kiara")
        kiara_msg = (
            f"CAMPAIGN VIDEO GENERATION REQUEST\n\n"
            f"Client: {client} | Platform: {platform}\n"
            f"Generate Arcads API job specs for ALL {len(video_days)} video slots.\n\n"
            f"DANTE'S SCRIPTS:\n{json.dumps(scripts, indent=2)}\n\n"
            f"CALENDAR VIDEO DAYS:\n{json.dumps(video_days, indent=2)}\n\n"
            f"For EACH video:\n"
            f"- Write a specific visual prompt (style, motion, mood, on-screen action)\n"
            f"- Output a complete ```json block: model, prompt, aspectRatio (9:16 for {platform}), duration (10-15s)\n\n"
            f"Default model: kling-3.0. One ```json block per video. No placeholders."
        )
        kiara_out = await call_agent(kiara_sys, kiara_msg, max_tokens=6000)

        # ── Step 5: Auto-fire videos to Arcads ───────────────────
        video_specs = extract_video_specs_from_kiara(kiara_out)
        video_job_ids: list = []
        if video_specs:
            set_step("generating_videos", f"Submitting {len(video_specs)} jobs to Arcads...")
            video_job_ids = await fire_video_jobs(video_specs, client, campaign_id)

        # ── Step 6: Assemble calendar ─────────────────────────────
        set_step("assembling", "Merging all outputs into final calendar...")
        cap_map = {c.get("day"): c for c in captions if isinstance(c, dict) and "day" in c}
        scr_map = {s.get("day"): s for s in scripts  if isinstance(s, dict) and "day" in s}

        for day in calendar_days:
            d = day.get("day")
            if d in cap_map:
                day["caption"]  = cap_map[d].get("caption", "")
                day["hashtags"] = cap_map[d].get("hashtags", [])
                day["cta"]      = cap_map[d].get("cta", "")
            if d in scr_map:
                day["hook"]             = scr_map[d].get("hook", "")
                day["script"]           = scr_map[d].get("script", "")
                day["on_screen_text"]   = scr_map[d].get("on_screen_text", "")
                day["visual_direction"] = scr_map[d].get("visual_direction", "")

        final = {
            "campaign_id":   campaign_id,
            "client":        client,
            "platform":      platform,
            "duration_days": duration_days,
            "style":         style_brief,
            "days":          calendar_days,
            "strategy_notes": cal_response,
            "video_job_ids": video_job_ids,
            "zara_research": zara_analysis,
            "generated_at":  datetime.now().isoformat(),
        }

        conn = sqlite3.connect(DB_PATH)
        conn.execute("""
            UPDATE campaigns
            SET status='complete', current_step='complete', calendar_json=?, video_job_ids=?
            WHERE id=?
        """, (json.dumps(final), json.dumps(video_job_ids), campaign_id))
        conn.commit()
        conn.close()

        _campaign_status[campaign_id] = {
            "step":    "complete",
            "message": f"{len(calendar_days)} days planned, {len(video_job_ids)} videos queued in Arcads",
            "calendar": final,
        }

    except Exception as exc:
        logging.error("Campaign %s failed: %s\n%s", campaign_id, exc, traceback.format_exc())
        try:
            conn = sqlite3.connect(DB_PATH)
            conn.execute("UPDATE campaigns SET status='failed', current_step=? WHERE id=?",
                         (str(exc)[:200], campaign_id))
            conn.commit()
            conn.close()
        except Exception:
            pass
        _campaign_status[campaign_id] = {"step": "failed", "message": str(exc)}
    finally:
        await asyncio.sleep(7200)
        _campaign_status.pop(campaign_id, None)


# ── GitHub + Anthropic ────────────────────────────────────────────

async def fetch_agent(name: str) -> str:
    if name in _agent_cache:
        return _agent_cache[name]
    url = f"https://raw.githubusercontent.com/{GITHUB_REPO}/main/agents/{name}.md"
    headers = {"Authorization": f"Bearer {GITHUB_TOKEN}"} if GITHUB_TOKEN else {}
    async with httpx.AsyncClient() as h:
        r = await h.get(url, headers=headers, timeout=15)
        r.raise_for_status()
    _agent_cache[name] = r.text
    return r.text

async def call_agent(system: str, message: str, model: str = "claude-sonnet-4-6",
                     max_tokens: int = 4096) -> str:
    r = await anthropic_client.messages.create(
        model=model, max_tokens=max_tokens, system=system,
        messages=[{"role": "user", "content": message}],
    )
    return r.content[0].text

# ── Save helpers ──────────────────────────────────────────────────

def _slug(text: str, n: int = 35) -> str:
    return re.sub(r"[^a-z0-9]+", "-", text[:n].lower()).strip("-")

def save_markdown(job: dict) -> Path:
    ts    = job["timestamp"]
    title = job.get("title") or job["brief"]
    path  = OUTPUTS_DIR / f"{ts.strftime('%Y-%m-%d_%H-%M')}_{_slug(title)}.md"
    client_label = {"wibiz": "WiBiz", "ai-living": "AI Living", "charter": "Club Charter"}.get(
        job.get("client", "other"), job.get("project_name") or "Other")
    lines = ["# Studio N — Job Output", "",
             f"**Title:** {title}", "",
             f"**Client:** {client_label}", "",
             f"**Date:** {ts.strftime('%Y-%m-%d %H:%M')}", "",
             f"**Brief:** {job['brief']}", "", "---", ""]
    if job.get("marcus_analysis"):
        lines += ["## Marcus — Brief Analysis", "", job["marcus_analysis"], "", "---", ""]
    for a, c in (job.get("stage1_outputs") or {}).items():
        lines += [f"## {a.capitalize()} — Stage 1", "", c, "", "---", ""]
    if job.get("marcus_review"):
        lines += ["## Marcus — Stage 1 Review", "", job["marcus_review"], "", "---", ""]
    for a, c in (job.get("cascade_outputs") or {}).items():
        lines += [f"## {a.capitalize()} — Cascade", "", c, "", "---", ""]
    if job.get("marcus_cascade_review"):
        lines += ["## Marcus — Final Review", "", job["marcus_cascade_review"], ""]
    path.write_text("\n".join(lines), encoding="utf-8")
    return path

def save_assembled(job: dict, assembly_type: str, content: str) -> Path:
    ts = job["timestamp"]
    ext, label = ASSEMBLERS[assembly_type]
    path = OUTPUTS_DIR / f"{ts.strftime('%Y-%m-%d_%H-%M')}_{_slug(job['brief'])}_{label}{ext}"
    path.write_text(content, encoding="utf-8")
    return path

# ── Markdown → HTML ───────────────────────────────────────────────

def _md_to_html(text: str) -> str:
    t = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    parts = t.split("**")
    t = "".join(f"<strong>{p}</strong>" if i % 2 else p for i, p in enumerate(parts))
    for pat, rep in [
        (r"^### (.+)$", r"<h3>\1</h3>"), (r"^## (.+)$", r"<h2>\1</h2>"),
        (r"^# (.+)$",   r"<h1>\1</h1>"), (r"^---+$",    "<hr>"),
        (r"^[-•] (.+)$", r"<li>\1</li>"),
    ]:
        t = re.sub(pat, rep, t, flags=re.MULTILINE)
    out = []
    for chunk in t.split("\n\n"):
        chunk = chunk.strip()
        if not chunk: continue
        if re.match(r"^<(h[1-6]|li|hr)", chunk):
            chunk = re.sub(r"(<li>.*</li>)", r"<ul>\1</ul>", chunk, flags=re.DOTALL)
            out.append(chunk)
        else:
            out.append(f"<p>{chunk.replace(chr(10), '<br>')}</p>")
    return "\n".join(out)

# ── Shared chrome ─────────────────────────────────────────────────

_PAGE_STYLE = """
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; background: #f8fafc; color: #1e293b; min-height: 100vh; }
  nav { background: #1e293b; padding: 14px 32px; display: flex; align-items: center; justify-content: space-between; }
  .nav-brand { color: white; font-size: 16px; font-weight: 700; text-decoration: none; letter-spacing: -0.3px; }
  .nav-brand span { color: #64748b; font-weight: 400; font-size: 13px; margin-left: 8px; }
  .nav-links { display: flex; gap: 6px; align-items: center; }
  .nav-link { color: #94a3b8; font-size: 13px; text-decoration: none; padding: 6px 12px; border-radius: 6px; transition: all 0.15s; }
  .nav-link:hover, .nav-link.active { background: rgba(255,255,255,0.1); color: white; }
  .container { max-width: 900px; margin: 0 auto; padding: 32px 24px 80px; }
  .dm-toggle { width:32px;height:32px;border-radius:7px;border:1px solid rgba(255,255,255,0.12);background:rgba(255,255,255,0.06);cursor:pointer;display:flex;align-items:center;justify-content:center;transition:all 0.15s;color:#94a3b8; }
  .dm-toggle:hover { background:rgba(255,255,255,0.14);color:white; }
  /* Job cards */
  .job-card { background:white;border-radius:12px;padding:20px 24px;margin-bottom:14px;box-shadow:0 1px 4px rgba(0,0,0,0.08); }
  .job-top { display:flex;align-items:flex-start;justify-content:space-between;gap:16px;margin-bottom:12px; }
  .job-title { font-size:15px;font-weight:700;color:#1e293b; }
  .job-ts { font-size:12px;color:#94a3b8;margin-bottom:6px;margin-top:3px; }
  .job-brief-text { font-size:12px;color:#64748b;line-height:1.5; }
  .job-agents { display:flex;flex-wrap:wrap;gap:4px;margin-bottom:10px; }
  .agent-chip { font-size:10px;font-weight:600;padding:2px 8px;border-radius:4px;background:#f1f5f9;color:#475569;white-space:nowrap; }
  .job-files { display:flex;flex-wrap:wrap;gap:6px;align-items:center; }
  .job-rebrief { font-size:11px;font-weight:700;padding:5px 12px;border-radius:6px;background:#f8fafc;color:#475569;text-decoration:none;border:1px solid #e2e8f0;white-space:nowrap;cursor:pointer;transition:all 0.15s; }
  .job-rebrief:hover { background:#0f172a;color:white;border-color:#0f172a; }
  /* Tabs */
  .tab-bar { display:flex;flex-wrap:wrap;gap:6px;margin-bottom:24px; }
  .tab-link { font-size:12px;font-weight:700;padding:6px 14px;border-radius:6px;text-decoration:none;white-space:nowrap;transition:all 0.15s; }
  /* Search */
  .search-bar { margin-bottom:20px; }
  .search-input { width:100%;border:1px solid #e2e8f0;border-radius:9px;padding:10px 14px;font-size:14px;font-family:inherit;color:#1e293b;outline:none;background:white;transition:border-color 0.15s; }
  .search-input:focus { border-color:#94a3b8; }
  /* Page head */
  .page-head { display:flex;align-items:baseline;justify-content:space-between;margin-bottom:16px; }
  .page-title { font-size:20px;font-weight:700; }
  .page-count { font-size:13px;color:#94a3b8; }
  /* Star button */
  .star-btn { background:none;border:none;cursor:pointer;font-size:16px;padding:0 2px;line-height:1;color:#cbd5e1;transition:color 0.15s;flex-shrink:0; }
  .star-btn:hover { color:#f59e0b; }
  .job-card.starred .star-btn { color:#f59e0b; }
  /* Checkbox */
  .job-check { width:15px;height:15px;cursor:pointer;flex-shrink:0;accent-color:#1e293b; }
  /* Bulk archive bar */
  .bulk-bar { display:none;position:sticky;bottom:20px;left:0;right:0;margin:0 auto;max-width:400px;
    background:#1e293b;color:white;border-radius:12px;padding:12px 20px;
    align-items:center;justify-content:space-between;gap:12px;
    box-shadow:0 4px 20px rgba(0,0,0,0.3);z-index:50; }
  .bulk-bar.visible { display:flex; }
  .bulk-bar-label { font-size:13px;font-weight:600; }
  .bulk-archive-btn { background:#ef4444;color:white;border:none;padding:7px 16px;border-radius:8px;
    font-size:12px;font-weight:700;cursor:pointer;font-family:inherit;transition:background 0.15s; }
  .bulk-archive-btn:hover { background:#dc2626; }
  .bulk-cancel-btn { background:rgba(255,255,255,0.1);color:white;border:none;padding:7px 12px;border-radius:8px;
    font-size:12px;font-weight:600;cursor:pointer;font-family:inherit;transition:background 0.15s; }
  .bulk-cancel-btn:hover { background:rgba(255,255,255,0.2); }
  /* ZIP button */
  .zip-btn { font-size:11px;font-weight:700;padding:5px 11px;border-radius:6px;
    background:#f0fdf4;color:#16a34a;text-decoration:none;white-space:nowrap;
    border:1px solid #86efac;transition:all 0.15s; }
  .zip-btn:hover { background:#dcfce7; }
  /* Mobile responsive */
  @media (max-width: 640px) {
    nav { padding: 12px 16px; }
    .container { padding: 16px 16px 60px; }
    .job-card { padding: 16px 18px; }
    .tab-bar { overflow-x: auto; flex-wrap: nowrap; -webkit-overflow-scrolling: touch; }
    .search-input { font-size: 16px; }
  }
  /* Dark mode */
  body.dark { background:#0b1120;color:#cbd5e1; }
  body.dark nav { background:#090f1c; }
  body.dark .page-title { color:#e2e8f0; }
  body.dark .page-count { color:#475569; }
  body.dark .job-card { background:#141e2e;box-shadow:0 1px 3px rgba(0,0,0,0.4); }
  body.dark .job-title { color:#e2e8f0; }
  body.dark .job-ts { color:#475569; }
  body.dark .job-brief-text { color:#475569; }
  body.dark .agent-chip { background:#1e293b;color:#64748b; }
  body.dark .job-rebrief { background:#1e293b;color:#64748b;border-color:#1e293b; }
  body.dark .job-rebrief:hover { background:#334155;color:#e2e8f0;border-color:#334155; }
  body.dark .search-input { background:#141e2e;border-color:#1e293b;color:#e2e8f0; }
  body.dark .search-input:focus { border-color:#334155; }
  body.dark .search-input::placeholder { color:#475569; }
"""

_PRINT_STYLE = """
  body { font-family: -apple-system, sans-serif; max-width: 780px; margin: 40px auto; padding: 0 24px 60px; color: #1e293b; }
  .toolbar { display:flex; gap:10px; margin-bottom:28px; }
  .btn { background:#1e293b; color:white; border:none; padding:9px 20px; border-radius:7px; font-size:13px; font-weight:600; cursor:pointer; }
  .btn-outline { background:none; border:1px solid #e2e8f0; color:#475569; }
  h1 { font-size:20px; font-weight:700; margin-bottom:4px; }
  .meta { color:#64748b; font-size:13px; margin-bottom:8px; }
  .brief-box { background:#f8fafc; border-left:3px solid #1e293b; padding:11px 16px; margin-bottom:28px; font-size:13px; color:#475569; border-radius:0 6px 6px 0; }
  h2{font-size:16px;font-weight:700;margin:28px 0 10px;} h3{font-size:14px;font-weight:700;margin:20px 0 8px;}
  p{font-size:14px;line-height:1.75;margin-bottom:12px;} ul{padding-left:20px;margin-bottom:12px;}
  li{font-size:14px;line-height:1.7;margin-bottom:4px;} strong{font-weight:700;}
  hr{border:none;border-top:1px solid #e2e8f0;margin:24px 0;}
  @media print { .toolbar{display:none!important;} body{margin:20px;} }
"""

def _nav(active: str) -> str:
    def lnk(href, label, key):
        cls = "nav-link active" if key == active else "nav-link"
        return f'<a href="{href}" class="{cls}">{label}</a>'
    return f"""<nav>
      <a href="/" class="nav-brand">Studio N<span>by Marcus</span></a>
      <div class="nav-links">
        {lnk('/','Brief','brief')}
        {lnk('/campaign','Campaign','campaign')}
        {lnk('/outputs','Outputs','outputs')}
        {lnk('/video-studio','Video Studio','studio')}
        <button class="dm-toggle" onclick="toggleDark()" title="Toggle dark mode">
          <svg id="dm-moon" width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M21 12.79A9 9 0 1 1 11.21 3 7 7 0 0 0 21 12.79z"/></svg>
          <svg id="dm-sun"  width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="display:none"><circle cx="12" cy="12" r="5"/><line x1="12" y1="1" x2="12" y2="3"/><line x1="12" y1="21" x2="12" y2="23"/><line x1="4.22" y1="4.22" x2="5.64" y2="5.64"/><line x1="18.36" y1="18.36" x2="19.78" y2="19.78"/><line x1="1" y1="12" x2="3" y2="12"/><line x1="21" y1="12" x2="23" y2="12"/><line x1="4.22" y1="19.78" x2="5.64" y2="18.36"/><line x1="18.36" y1="5.64" x2="19.78" y2="4.22"/></svg>
        </button>
        <a href="/logout" class="nav-link" style="color:#6b7280;">Sign out</a>
      </div>
    </nav>
    <script>
    (function(){{if(localStorage.getItem('dm')==='1'){{document.body.classList.add('dark');var m=document.getElementById('dm-moon'),s=document.getElementById('dm-sun');if(m)m.style.display='none';if(s)s.style.display='';}} }})();
    function toggleDark(){{var d=document.body.classList.toggle('dark');localStorage.setItem('dm',d?'1':'0');var m=document.getElementById('dm-moon'),s=document.getElementById('dm-sun');if(m)m.style.display=d?'none':'';if(s)s.style.display=d?'':'none';}}
    </script>"""

def _print_page(title: str, brief: str, date: str, body: str) -> str:
    return f"""<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><title>{title}</title>
<style>{_PRINT_STYLE}</style></head><body>
<div class="toolbar">
  <button class="btn" onclick="window.print()">Save as PDF</button>
  <button class="btn btn-outline" onclick="window.close()">Close</button>
</div>
<h1>{title}</h1><div class="meta">{date}</div>
<div class="brief-box"><strong>Brief:</strong> {brief}</div>
{body}</body></html>"""

# ── Agent extraction (lightweight second call) ────────────────────

async def extract_agents_from_analysis(analysis: str, brief: str) -> list[str]:
    """Dedicated extraction call — returns only a JSON array of agent names.
    Uses a cheap model with max_tokens=200 so it has nothing to do but output the array."""
    prompt = (
        f"Marketing brief: {brief}\n\n"
        f"Strategic analysis:\n{analysis[:2000]}\n\n"
        "Which specialist agents are needed for this project? "
        "Choose from ONLY these exact names: callum, priya, dante, suki, felix, nadia, zara, reeva, kiara, rex, nova\n"
        "dante = video scripts and shot lists. kiara = AI video generation. rex = programmatic video code. nova = brand visuals.\n"
        "Return only a JSON array of names from the list above. Example: [\"dante\", \"priya\"]\n"
        "Do NOT invent names not on this list. No explanation. No other text. Just the array."
    )
    try:
        raw = await call_agent(
            "You extract agent names from marketing briefs. Output only a JSON array of names, nothing else.",
            prompt,
            model="claude-haiku-4-5-20251001",
            max_tokens=200,
        )
        m = re.search(r"\[.*?\]", raw, re.DOTALL)
        if m:
            names = json.loads(m.group(0))
            valid = [n for n in names if n in VALID_AGENTS]
            logging.info("extraction call returned %s → valid agents: %s", names, valid)
            return valid
        logging.warning("extraction call had no JSON array in response: %r", raw[:200])
    except Exception as e:
        logging.error("extraction call failed: %s", e)
    return []


# ── Job processor ─────────────────────────────────────────────────

async def process_job(job_id: str, brief: str, title: str = "",
                      client: str = "other", project_name: str = "",
                      allowed_agents: list[str] | None = None) -> None:
    q = _jobs[job_id]

    async def emit(ev: dict) -> None:
        if ev.get("type") == "status":
            _processing_status[job_id] = ev.get("message", "")
        await q.put(ev)

    job: dict = {
        "brief": brief,
        "timestamp": datetime.now(),
        "title": title or brief[:60],
        "client": client,
        "project_name": project_name,
    }

    # Write a processing stub immediately — job appears in Outputs right away.
    db_save_job_initial(job_id, title or brief[:60], client, project_name, brief)

    try:
        # Build context-enriched brief — tells Marcus and all agents exactly which
        # client or project this is for, so they don't have to guess from tone clues.
        _CLIENT_LABELS = {"wibiz": "WiBiz", "ai-living": "AI Living", "charter": "Club Charter"}
        _label = _CLIENT_LABELS.get(client)
        if _label:
            enriched_brief = f"CLIENT: {_label}\n\nBRIEF:\n{brief}"
        elif project_name:
            enriched_brief = f"PROJECT: {project_name}\n\nBRIEF:\n{brief}"
        else:
            enriched_brief = brief

        await emit({"type": "status", "message": "Marcus is reading your brief..."})
        marcus_system = await fetch_agent("marcus")
        # Prefix puts the JSON requirement as the very first thing Marcus reads,
        # before his own identity prompt, so it can't be overridden by prose flow.
        marcus_raw = await call_agent(
            ORCHESTRATION_PREFIX + marcus_system, enriched_brief, model="claude-opus-4-6")

        # Strip JSON block from analysis text if Marcus included one.
        json_match      = re.search(r"```json\s*(\{.*?\})\s*```", marcus_raw, re.DOTALL)
        marcus_analysis = marcus_raw[:json_match.start()].strip() if json_match else marcus_raw
        agents_needed:  list[str] = []
        agent_briefs:   dict[str, str] = {}

        # Layer 1: try to parse Marcus's own JSON block (best quality — includes written briefs).
        if json_match:
            try:
                p             = json.loads(json_match.group(1))
                agents_needed = p.get("agents_needed", [])
                agent_briefs  = p.get("briefs", {})
                logging.info("job %s — layer1: parsed Marcus JSON block, agents=%s", job_id, agents_needed)
            except json.JSONDecodeError as e:
                logging.error("job %s — layer1: Marcus JSON parse failed: %s", job_id, e)
        else:
            logging.warning("job %s — layer1: no JSON block in Marcus response", job_id)

        # Layer 2: dedicated extraction call — lightweight model, max_tokens=200, returns only an array.
        # Runs whenever layer 1 produced no agents. Cannot be confused by prose.
        if not agents_needed:
            await emit({"type": "status", "message": "Extracting agent plan..."})
            agents_needed = await extract_agents_from_analysis(marcus_analysis, enriched_brief)
            if agents_needed:
                agent_briefs = {n: f"Original brief: {enriched_brief}\n\nMarcus's analysis:\n{marcus_analysis}"
                                for n in agents_needed}
                logging.info("job %s — layer2: extraction call resolved agents=%s", job_id, agents_needed)
            else:
                logging.warning("job %s — layer2: extraction call returned no agents", job_id)

        # Layer 3: last-resort text scan for agent name mentions.
        if not agents_needed:
            mentioned = [n for n in VALID_AGENTS if re.search(rf"\b{n}\b", marcus_raw, re.IGNORECASE)]
            if mentioned:
                agents_needed = mentioned
                agent_briefs  = {n: f"Original brief: {enriched_brief}\n\nMarcus's analysis:\n{marcus_analysis}"
                                 for n in mentioned}
                logging.warning("job %s — layer3: text-scan inferred agents=%s", job_id, mentioned)
            else:
                logging.warning("job %s — layer3: no agents found anywhere; pipeline will produce no output", job_id)

        logging.info("job %s — final agents_needed: %s", job_id, agents_needed)
        job["marcus_analysis"] = marcus_analysis
        await emit({"type": "marcus_analysis", "content": marcus_analysis})

        valid_s1 = [n for n in agents_needed if n in agent_briefs and n in VALID_AGENTS]
        if allowed_agents:
            valid_s1 = [n for n in valid_s1 if n in allowed_agents]
        if valid_s1:
            await emit({"type": "agents_identified", "agents": valid_s1})
            await emit({"type": "status",
                        "message": f"Running {', '.join(n.capitalize() for n in valid_s1)} in parallel..."})

        async def run_s1(name: str) -> tuple[str, str]:
            sys_ = await fetch_agent(name)
            out  = await call_agent(sys_, agent_briefs[name])
            return name, out

        stage1_outputs: dict[str, str] = {}
        if valid_s1:
            results = await asyncio.gather(*[run_s1(n) for n in valid_s1])
            stage1_outputs = dict(results)
            # Zara: append live competitor Instagram data when brief signals research
            if "zara" in stage1_outputs:
                await emit({"type": "status", "message": "Zara: running competitor intelligence scrape..."})
                stage1_outputs["zara"] = await enrich_zara_with_apify(stage1_outputs["zara"], enriched_brief)
            for name, content in stage1_outputs.items():
                logging.info("job %s — stage1 %s: %d chars", job_id, name, len(content))
                await emit({"type": "specialist", "agent": name, "content": content})
        else:
            logging.warning("job %s — no valid stage1 agents to run (valid_s1=%s)", job_id, valid_s1)

        job["stage1_outputs"] = stage1_outputs

        # Stage 1 review + cascade decision
        marcus_review = None
        cascade_plan  = {"next_agents": [], "assembly": "none", "flags": []}

        if stage1_outputs:
            await emit({"type": "status", "message": "Marcus is reviewing outputs..."})
            review_input = (
                "Review each output against the brief and brand standards. "
                "Give your verdict (Approved / Revise / Reject) for each.\n\n"
                + "".join(f"## {n.capitalize()}\n\n{o}\n\n---\n\n"
                          for n, o in stage1_outputs.items())
            )
            review_raw = await call_agent(
                marcus_system + REVIEW_CASCADE_SUFFIX,
                review_input, model="claude-opus-4-6")

            cm = re.search(r"```cascade\s*(\{.*?\})\s*```", review_raw, re.DOTALL)
            marcus_review = review_raw[:cm.start()].strip() if cm else review_raw

            if cm:
                try:
                    cascade_plan = json.loads(cm.group(1)).get("cascade", cascade_plan)
                except json.JSONDecodeError as e:
                    logging.error("job %s — failed to parse Marcus cascade JSON: %s\nRaw block: %s",
                                  job_id, e, cm.group(1)[:500])
            else:
                logging.warning("job %s — Marcus review returned no cascade block", job_id)

            logging.info("job %s — cascade plan: %s", job_id, cascade_plan)
            job["marcus_review"] = marcus_review
            await emit({"type": "review", "content": marcus_review})

        # Stage 2: cascade agents
        cascade_outputs: dict[str, str] = {}
        next_agents = [n for n in cascade_plan.get("next_agents", [])
                       if n in VALID_AGENTS and n not in stage1_outputs]
        if allowed_agents:
            next_agents = [n for n in next_agents if n in allowed_agents]

        if next_agents:
            await emit({"type": "cascade_start", "agents": next_agents})
            await emit({"type": "status",
                        "message": f"Cascade: {', '.join(n.capitalize() for n in next_agents)}..."})

            ctx = "\n\n".join(
                f"## {n.capitalize()} — Approved\n\n{o}" for n, o in stage1_outputs.items())
            base_brief = (
                f"Original brief: {enriched_brief}\n\n---\n\n"
                f"APPROVED OUTPUTS FROM PREVIOUS STAGE (use as source of truth):\n\n{ctx}\n\n---\n\n"
                f"Produce complete, final content — not a draft. This goes directly into the assembled deliverable."
            )

            async def run_cascade(name: str) -> tuple[str, str]:
                sys_ = await fetch_agent(name)
                out  = await call_agent(sys_, base_brief)
                return name, out

            results = await asyncio.gather(*[run_cascade(n) for n in next_agents])
            cascade_outputs = dict(results)
            for name, content in cascade_outputs.items():
                logging.info("job %s — cascade %s: %d chars", job_id, name, len(content))
                await emit({"type": "cascade_output", "agent": name, "content": content})

        job["cascade_outputs"] = cascade_outputs

        for flag in cascade_plan.get("flags", []):
            msg = THIRD_PARTY_MESSAGES.get(flag, f"Manual step: {flag}")
            await emit({"type": "third_party_flag", "flag": flag, "message": msg})

        if cascade_outputs:
            await emit({"type": "status", "message": "Marcus is reviewing cascade outputs..."})
            r2 = await call_agent(
                marcus_system,
                "Review these cascade outputs:\n\n" +
                "".join(f"## {n.capitalize()}\n\n{o}\n\n---\n\n"
                        for n, o in cascade_outputs.items()),
                model="claude-opus-4-6")
            job["marcus_cascade_review"] = r2
            await emit({"type": "cascade_review", "content": r2})

        # Assembly — generate and save each file sequentially so one failure
        # never blocks the others. Each step logs on error but continues.
        all_outputs = {**stage1_outputs, **cascade_outputs}
        assembled_content_for_db: dict[str, str] = {}

        logging.info(
            "job %s — all_outputs before assembly: keys=%s stage1=%s cascade=%s total_chars=%d",
            job_id,
            list(all_outputs.keys()),
            list(stage1_outputs.keys()),
            list(cascade_outputs.keys()),
            sum(len(v) for v in all_outputs.values()),
        )
        if not all_outputs:
            logging.warning("job %s — all_outputs is empty, skipping assembly entirely", job_id)

        if all_outputs:
            await emit({"type": "assembly_start", "assembly_type": "bundle"})
            bundle = {}

            # 1. HTML Website
            await emit({"type": "status", "message": "Building HTML website..."})
            try:
                website_content = await assemble_html(
                    _build_website_prompt(all_outputs, brief),
                    "You generate complete, production-ready B2B HTML websites with Google Fonts. Output only valid HTML.")
                website_path = save_assembled(job, "html_website", website_content)
                bundle["html_website"] = website_path.name
                assembled_content_for_db["html_website"] = website_content
                logging.info("Assembly saved: %s", website_path)
            except Exception as e:
                logging.error("Assembly html_website failed: %s\n%s", e, traceback.format_exc())

            # 2. A4 One-Pager
            await emit({"type": "status", "message": "Building A4 one-pager..."})
            try:
                onepager_content = await assemble_html(
                    _build_onepager_prompt(all_outputs, brief),
                    "You generate complete, print-ready A4 HTML one-pagers with Google Fonts. Output only valid HTML.")
                onepager_path = save_assembled(job, "html_onepager", onepager_content)
                bundle["html_onepager"] = onepager_path.name
                assembled_content_for_db["html_onepager"] = onepager_content
                logging.info("Assembly saved: %s", onepager_path)
            except Exception as e:
                logging.error("Assembly html_onepager failed: %s\n%s", e, traceback.format_exc())

            # 3. Canva JSON Template
            await emit({"type": "status", "message": "Building Canva JSON template..."})
            try:
                canva_json_str = await assemble_canva_json(all_outputs, brief)
                canva_path = save_assembled(job, "canva_json", canva_json_str)
                bundle["canva_json"] = canva_path.name
                assembled_content_for_db["canva_json"] = canva_json_str
                logging.info("Assembly saved: %s", canva_path)
            except Exception as e:
                logging.error("Assembly canva_json failed: %s\n%s", e, traceback.format_exc())

            # 4. Social Content Pack (Callum + Priya in parallel)
            await emit({"type": "status", "message": "Generating social content pack (Callum + Priya)..."})
            try:
                social_cascade_content = await assemble_social_cascade(all_outputs, brief)
                social_cascade_path = save_assembled(job, "social_pack_cascade", social_cascade_content)
                bundle["social_pack_cascade"] = social_cascade_path.name
                assembled_content_for_db["social_pack_cascade"] = social_cascade_content
                logging.info("Assembly saved: %s", social_cascade_path)
                await post_to_buffer(social_cascade_content)
            except Exception as e:
                logging.error("Assembly social_pack_cascade failed: %s\n%s", e, traceback.format_exc())

            # 5. Video Brief (Dante)
            await emit({"type": "status", "message": "Generating video brief (Dante)..."})
            try:
                video_brief_content = await assemble_video_brief(all_outputs, brief)
                video_brief_path = save_assembled(job, "video_brief", video_brief_content)
                bundle["video_brief"] = video_brief_path.name
                assembled_content_for_db["video_brief"] = video_brief_content
                logging.info("Assembly saved: %s", video_brief_path)
            except Exception as e:
                logging.error("Assembly video_brief failed: %s\n%s", e, traceback.format_exc())

            await emit({
                "type":  "assembly_bundle_done",
                "files": bundle,
                "labels": {
                    "html_website":        "HTML Website",
                    "html_onepager":       "A4 One-Pager (PDF-ready)",
                    "canva_json":          "Canva Template (JSON)",
                    "social_pack_cascade": "Social Content Pack",
                    "video_brief":         "Video Brief (PDF-ready)",
                },
            })

            # Optional compiled social pack (legacy social_pack assembly type)
            if cascade_plan.get("assembly") == "social_pack":
                await emit({"type": "status", "message": "Compiling full social content pack..."})
                try:
                    social_content = await assemble_social_pack(all_outputs, brief)
                    social_path = save_assembled(job, "social_pack", social_content)
                    await emit({
                        "type":          "assembly_done",
                        "assembly_type": "social_pack",
                        "label":         "Social Content Pack",
                        "filename":      social_path.name,
                    })
                except Exception as e:
                    logging.error("Assembly social_pack failed: %s", e)

        job["assembled_files"] = bundle if all_outputs else {}
        md_path = save_markdown(job)
        db_save_job(job_id, job, assembled_content_for_db)
        _completed[job_id] = job
        await emit({"type": "done", "job_id": job_id, "saved_as": md_path.name})

    except Exception as exc:
        logging.error("process_job %s failed: %s\n%s", job_id, exc, traceback.format_exc())
        db_update_job_status(job_id, "failed")
        _processing_status[job_id] = "__failed__"
        await emit({"type": "error", "message": str(exc)})
    finally:
        await asyncio.sleep(7200)
        _jobs.pop(job_id, None)
        _completed.pop(job_id, None)
        _processing_status.pop(job_id, None)

# ── Routes ────────────────────────────────────────────────────────

@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request):
    if request.session.get("authenticated"):
        return RedirectResponse("/", status_code=302)
    return templates.TemplateResponse(request=request, name="login.html", context={"error": None})

@app.post("/login", response_class=HTMLResponse)
async def login_submit(request: Request):
    form = await request.form()
    username = form.get("username", "")
    password = form.get("password", "")
    if (secrets.compare_digest(username, HTTP_USER) and
            secrets.compare_digest(password, HTTP_PASS)):
        request.session["authenticated"] = True
        return RedirectResponse("/", status_code=302)
    return templates.TemplateResponse(request=request, name="login.html", context={"error": "Incorrect username or password."})

@app.get("/logout")
async def logout(request: Request):
    request.session.clear()
    return RedirectResponse("/login", status_code=302)

@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    return templates.TemplateResponse(request=request, name="index.html")

async def parse_uploaded_file(filename: str, data: bytes) -> str:
    """Extract text content from an uploaded file. Images are described via Claude vision."""
    import io
    ext = Path(filename).suffix.lower()

    if ext in ('.txt', '.md'):
        return data.decode('utf-8', errors='replace')

    elif ext == '.pdf':
        try:
            import pdfplumber
            parts = []
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        parts.append(t)
            return '\n\n'.join(parts) if parts else '[PDF: no extractable text]'
        except Exception as e:
            return f'[PDF parse error: {e}]'

    elif ext == '.docx':
        try:
            import docx
            doc = docx.Document(io.BytesIO(data))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    paragraphs.append('\t'.join(c.text for c in row.cells if c.text.strip()))
            return '\n'.join(paragraphs)
        except Exception as e:
            return f'[DOCX parse error: {e}]'

    elif ext in ('.xlsx', '.xls'):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            parts = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                parts.append(f'Sheet: {sheet_name}')
                for row in ws.iter_rows(values_only=True):
                    row_text = '\t'.join(str(c) if c is not None else '' for c in row)
                    if row_text.strip():
                        parts.append(row_text)
            return '\n'.join(parts)
        except Exception as e:
            return f'[Spreadsheet parse error: {e}]'

    elif ext == '.csv':
        try:
            import csv
            reader = csv.reader(io.StringIO(data.decode('utf-8', errors='replace')))
            return '\n'.join('\t'.join(row) for row in reader)
        except Exception as e:
            return f'[CSV parse error: {e}]'

    elif ext in ('.png', '.jpg', '.jpeg', '.gif', '.webp'):
        try:
            import base64
            media_map = {'.png': 'image/png', '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
                         '.gif': 'image/gif', '.webp': 'image/webp'}
            media_type = media_map.get(ext, 'image/jpeg')
            b64 = base64.standard_b64encode(data).decode()
            resp = await anthropic_client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=1500,
                messages=[{"role": "user", "content": [
                    {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": b64}},
                    {"type": "text", "text": "Describe all content in this image in detail. Extract any text verbatim. Include layout, structure, and key visual information."}
                ]}]
            )
            return resp.content[0].text
        except Exception as e:
            return f'[Image OCR error: {e}]'

    else:
        return f'[Unsupported file type: {ext}]'


@app.post("/api/brief")
async def start_brief(
    brief:          str  = Form(...),
    title:          str  = Form(...),
    client_name:    str  = Form("other", alias="client"),
    project_name:   str  = Form(""),
    allowed_agents: str  = Form(""),
    files:          List[UploadFile] = File(default=[]),
):
    brief        = brief.strip()
    title        = title.strip()
    client_name  = client_name.strip()
    project_name = project_name.strip()

    if not brief:
        return JSONResponse({"error": "Brief is empty"}, status_code=400)
    if not title:
        return JSONResponse({"error": "Title is required"}, status_code=400)

    allowed_list: list[str] | None = None
    if allowed_agents.strip():
        allowed_list = [a.strip() for a in allowed_agents.split(",") if a.strip() in VALID_AGENTS]
        if not allowed_list:
            allowed_list = None

    # Parse uploaded files and append extracted content to the brief
    if files:
        attachments = []
        for f in files:
            if f.filename:
                data = await f.read()
                extracted = await parse_uploaded_file(f.filename, data)
                attachments.append(f"--- Attached file: {f.filename} ---\n{extracted}")
                logging.info("Parsed attachment: %s (%d bytes)", f.filename, len(data))
        if attachments:
            brief = brief + "\n\n" + "\n\n".join(attachments)

    job_id = str(uuid.uuid4())
    _jobs[job_id] = asyncio.Queue()
    asyncio.create_task(process_job(job_id, brief, title=title, client=client_name,
                                    project_name=project_name, allowed_agents=allowed_list))
    return {"job_id": job_id}

@app.post("/api/job/{job_id}/archive")
async def archive_job(job_id: str):
    conn = sqlite3.connect(DB_PATH)
    conn.execute("UPDATE jobs SET archived=1 WHERE id=?", (job_id,))
    conn.commit()
    conn.close()
    return {"ok": True}

@app.post("/api/job/{job_id}/unarchive")
async def unarchive_job(job_id: str):
    conn = sqlite3.connect(DB_PATH)
    conn.execute("UPDATE jobs SET archived=0 WHERE id=?", (job_id,))
    conn.commit()
    conn.close()
    return {"ok": True}

@app.get("/outputs/zip/{job_id}")
async def zip_job(job_id: str):
    import io
    import zipfile
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    row = conn.execute("SELECT title, assembled_files FROM jobs WHERE id=?", (job_id,)).fetchone()
    conn.close()
    if not row:
        return JSONResponse({"error": "not found"}, status_code=404)
    title = row["title"] or "job"
    try:
        af = json.loads(row["assembled_files"] or "{}")
    except Exception:
        af = {}
    safe_title = re.sub(r"[^a-z0-9]+", "-", title[:40].lower()).strip("-") or "job"
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for key, fname in af.items():
            path = OUTPUTS_DIR / fname
            if path.exists():
                zf.write(path, fname)
    buf.seek(0)
    return StreamingResponse(
        iter([buf.read()]),
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{safe_title}.zip"'},
    )

@app.get("/api/job/{job_id}")
async def get_job(job_id: str):
    conn = sqlite3.connect(DB_PATH)
    row = conn.execute(
        "SELECT title, brief, client, project_name FROM jobs WHERE id=?", (job_id,)
    ).fetchone()
    conn.close()
    if not row:
        return JSONResponse({"error": "not found"}, status_code=404)
    title, brief, client_val, project_name = row
    # Strip file attachment blocks — user just wants the original brief text
    if "\n\n--- Attached file:" in brief:
        brief = brief.split("\n\n--- Attached file:")[0]
    return {"title": title, "brief": brief, "client": client_val, "project_name": project_name or ""}

@app.get("/api/job/{job_id}/progress")
async def job_progress(job_id: str):
    """Poll endpoint for the Outputs page — returns current processing state."""
    if job_id in _completed:
        return {"status": "complete", "message": "Done"}
    if job_id in _jobs:
        msg = _processing_status.get(job_id, "Generating…")
        if msg == "__failed__":
            return {"status": "failed", "message": "Generation failed"}
        return {"status": "processing", "message": msg}
    # Check DB for persisted status
    conn = sqlite3.connect(DB_PATH)
    row = conn.execute("SELECT status FROM jobs WHERE id=?", (job_id,)).fetchone()
    conn.close()
    if row:
        return {"status": row[0] or "complete", "message": ""}
    return JSONResponse({"error": "not found"}, status_code=404)

@app.get("/api/stream/{job_id}")
async def stream_job(job_id: str):
    if job_id not in _jobs:
        return JSONResponse({"error": "Job not found"}, status_code=404)
    async def generate():
        q = _jobs[job_id]
        while True:
            try:
                event = await asyncio.wait_for(q.get(), timeout=15)
            except asyncio.TimeoutError:
                yield ": keepalive\n\n"
                continue
            yield f"data: {json.dumps(event)}\n\n"
            if event["type"] in ("done", "error"):
                break
    return StreamingResponse(generate(), media_type="text/event-stream",
                             headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


# ── Export routes ─────────────────────────────────────────────────

def _load_job_for_export(job_id: str) -> dict | None:
    """Return job dict from in-memory cache or SQLite fallback."""
    job = _completed.get(job_id)
    if job:
        return job
    conn = sqlite3.connect(DB_PATH)
    row = conn.execute(
        "SELECT brief, timestamp, marcus_analysis, stage1_outputs, marcus_review, "
        "cascade_outputs, marcus_cascade_review FROM jobs WHERE id=? AND status='complete'",
        (job_id,)
    ).fetchone()
    conn.close()
    if not row:
        return None
    brief, ts_str, marcus_analysis, stage1_json, marcus_review, cascade_json, marcus_cascade_review = row
    try:
        ts = datetime.fromisoformat(ts_str)
    except Exception:
        ts = datetime.utcnow()
    return {
        "brief": brief or "",
        "timestamp": ts,
        "marcus_analysis": marcus_analysis or "",
        "stage1_outputs": json.loads(stage1_json or "{}"),
        "marcus_review": marcus_review or "",
        "cascade_outputs": json.loads(cascade_json or "{}"),
        "marcus_cascade_review": marcus_cascade_review or "",
    }

@app.get("/api/export/{job_id}/{agent}.md")
async def export_md(job_id: str, agent: str):
    job = _load_job_for_export(job_id)
    if not job:
        return JSONResponse({"error": "Output not found or expired"}, status_code=404)
    ts = job["timestamp"]
    all_out = {**job.get("stage1_outputs",{}), **job.get("cascade_outputs",{})}
    if agent == "full":
        path = save_markdown(job); content = path.read_text(); fname = path.name
    elif agent == "review":
        content = job.get("marcus_review") or job.get("marcus_cascade_review") or ""
        fname = f"marcus-review_{ts.strftime('%Y-%m-%d_%H-%M')}.md"
    elif agent == "analysis":
        content = job.get("marcus_analysis", "")
        fname = f"marcus-analysis_{ts.strftime('%Y-%m-%d_%H-%M')}.md"
    else:
        content = all_out.get(agent, "")
        if not content:
            return JSONResponse({"error": "Not found"}, status_code=404)
        fname = f"{agent}_{ts.strftime('%Y-%m-%d_%H-%M')}.md"
    return StreamingResponse(iter([content.encode()]), media_type="text/markdown",
                             headers={"Content-Disposition": f'attachment; filename="{fname}"'})

@app.get("/api/export/{job_id}/{agent}/print", response_class=HTMLResponse)
async def print_export(job_id: str, agent: str):
    job = _load_job_for_export(job_id)
    if not job:
        return HTMLResponse("<h1>Output not found or expired</h1>", status_code=404)
    ts = job["timestamp"].strftime("%Y-%m-%d %H:%M")
    all_out = {**job.get("stage1_outputs",{}), **job.get("cascade_outputs",{})}
    if agent == "full":
        title, md_body = "Full Output", ""
        for s, k in [("Brief Analysis","marcus_analysis"),("Stage 1 Review","marcus_review"),
                     ("Final Review","marcus_cascade_review")]:
            if job.get(k): md_body += f"## Marcus — {s}\n\n{job[k]}\n\n---\n\n"
        for n, o in all_out.items(): md_body += f"## {n.capitalize()}\n\n{o}\n\n---\n\n"
    elif agent in ("review","analysis"):
        k = "marcus_review" if agent == "review" else "marcus_analysis"
        title, md_body = f"Marcus — {agent.capitalize()}", job.get(k,"")
    else:
        title, md_body = f"{agent.capitalize()} — Output", all_out.get(agent,"")
    return HTMLResponse(_print_page(title, job["brief"], ts, _md_to_html(md_body)))

# ── Studio page — redirect to /outputs ───────────────────────────

@app.get("/studio")
async def studio_redirect():
    return RedirectResponse("/outputs", status_code=302)

@app.get("/studio_old", response_class=HTMLResponse)
async def studio_page():
    html_files = sorted(OUTPUTS_DIR.glob("*.html"),
                        key=lambda f: f.stat().st_mtime, reverse=True)

    if not html_files:
        cards_html = """<div style="text-align:center;padding:80px 0;color:#94a3b8;">
          <div style="font-size:40px;margin-bottom:16px;">🎨</div>
          <div style="font-size:16px;font-weight:600;color:#475569;margin-bottom:8px;">No HTML outputs yet</div>
          <div style="font-size:13px;">Submit a brand, website, deck, or one-pager brief to generate your first visual output.</div>
        </div>"""
    else:
        cards_html = '<div class="studio-grid">'
        for f in html_files:
            size_kb = round(f.stat().st_size / 1024, 1)
            mtime   = datetime.fromtimestamp(f.stat().st_mtime).strftime("%Y-%m-%d %H:%M")
            label   = next((l for l in ["website","one-pager","deck"] if l in f.name), "html")
            label_colors = {"website": "#dbeafe:#1d4ed8",
                            "one-pager": "#d1fae5:#065f46",
                            "deck": "#ede9fe:#6d28d9"}
            lbg, lfg = label_colors.get(label, "#f1f5f9:#475569").split(":")
            cards_html += f"""
            <div class="studio-card">
              <div class="preview-wrap">
                <iframe src="/outputs/view/{f.name}" sandbox="allow-same-origin"
                        scrolling="no" class="preview-frame"></iframe>
                <div class="preview-overlay">
                  <a href="/outputs/view/{f.name}" target="_blank" class="overlay-btn">Open →</a>
                </div>
              </div>
              <div class="card-meta">
                <div class="card-top">
                  <span class="label-badge" style="background:{lbg};color:{lfg};">{label.upper()}</span>
                  <span class="card-size">{size_kb} KB</span>
                </div>
                <div class="card-name" title="{f.name}">{f.name}</div>
                <div class="card-date">{mtime}</div>
                <div class="card-actions">
                  <a href="/outputs/view/{f.name}" target="_blank" class="ca-btn ca-view">Open</a>
                  <a href="/outputs/download/{f.name}" class="ca-btn ca-dl">⬇ Download</a>
                </div>
              </div>
            </div>"""
        cards_html += "</div>"

    count = len(html_files)
    return HTMLResponse(f"""<!DOCTYPE html><html lang="en"><head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>Studio — Studio N</title>
<style>
  {_PAGE_STYLE}
  .page-head {{display:flex;align-items:baseline;justify-content:space-between;margin-bottom:24px;}}
  .page-title {{font-size:20px;font-weight:700;}}
  .page-count {{font-size:13px;color:#94a3b8;}}
  .studio-grid {{display:grid;grid-template-columns:repeat(auto-fill,minmax(280px,1fr));gap:20px;}}
  .studio-card {{background:white;border-radius:12px;overflow:hidden;box-shadow:0 1px 3px rgba(0,0,0,0.08);transition:box-shadow 0.15s;}}
  .studio-card:hover {{box-shadow:0 4px 16px rgba(0,0,0,0.12);}}
  .preview-wrap {{position:relative;height:200px;overflow:hidden;background:#f1f5f9;cursor:pointer;}}
  .preview-frame {{width:200%;height:200%;transform:scale(0.5);transform-origin:0 0;border:none;pointer-events:none;}}
  .preview-overlay {{position:absolute;inset:0;background:rgba(30,41,59,0);display:flex;align-items:center;
                     justify-content:center;transition:background 0.2s;}}
  .preview-wrap:hover .preview-overlay {{background:rgba(30,41,59,0.5);}}
  .overlay-btn {{opacity:0;background:white;color:#1e293b;padding:8px 18px;border-radius:8px;
                 font-size:13px;font-weight:700;text-decoration:none;transition:opacity 0.2s;}}
  .preview-wrap:hover .overlay-btn {{opacity:1;}}
  .card-meta {{padding:14px 16px;}}
  .card-top {{display:flex;align-items:center;justify-content:space-between;margin-bottom:8px;}}
  .label-badge {{font-size:10px;font-weight:700;padding:2px 8px;border-radius:4px;letter-spacing:0.5px;}}
  .card-size {{font-size:11px;color:#94a3b8;}}
  .card-name {{font-size:12px;font-weight:600;color:#1e293b;margin-bottom:3px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;}}
  .card-date {{font-size:11px;color:#94a3b8;margin-bottom:10px;}}
  .card-actions {{display:flex;gap:6px;}}
  .ca-btn {{font-size:11px;font-weight:600;padding:5px 10px;border-radius:6px;text-decoration:none;transition:all 0.15s;border:1px solid #e2e8f0;color:#475569;}}
  .ca-btn:hover {{background:#f8fafc;color:#1e293b;}}
  .ca-view {{border-color:#bbf7d0;color:#16a34a;}}
  .ca-view:hover {{background:#f0fdf4;}}
  .ca-dl {{border-color:#bfdbfe;color:#2563eb;}}
  .ca-dl:hover {{background:#eff6ff;}}
</style></head><body>
{_nav('studio')}
<div class="container">
  <div class="page-head">
    <span class="page-title">Studio</span>
    <span class="page-count">{count} HTML file{'s' if count!=1 else ''}</span>
  </div>
  {cards_html}
</div></body></html>""")

# ── Outputs page (DB-backed) ──────────────────────────────────────

CLIENT_META = {
    "wibiz":     {"label": "WiBiz",       "bg": "#dcfce7", "fg": "#15803d"},
    "ai-living": {"label": "AI Living",   "bg": "#dbeafe", "fg": "#1d4ed8"},
    "charter":   {"label": "Club Charter","bg": "#fef9c3", "fg": "#a16207"},
    "other":     {"label": "Other",       "bg": "#f1f5f9", "fg": "#475569"},
}

FILE_META = {
    "html_website":        {"label": "Website",      "bg": "#dbeafe", "fg": "#1d4ed8"},
    "html_onepager":       {"label": "One-Pager",    "bg": "#d1fae5", "fg": "#065f46"},
    "canva_json":          {"label": "Canva JSON",   "bg": "#ede9fe", "fg": "#6d28d9"},
    "social_pack_cascade": {"label": "Social Pack",  "bg": "#fef3c7", "fg": "#92400e"},
    "video_brief":         {"label": "Video Brief",  "bg": "#fee2e2", "fg": "#b91c1c"},
    "social_pack":         {"label": "Social Pack",  "bg": "#fef3c7", "fg": "#92400e"},
}

AGENT_ROLES = {
    "marcus": "Director of Marketing",
    "reeva":  "Brand & Identity",
    "callum": "LinkedIn & Long-Form",
    "priya":  "Social & Short-Form",
    "dante":  "Video & Reels",
    "suki":   "Static & Visual",
    "felix":  "Decks & Presentations",
    "nadia":  "Web & Copy",
    "zara":   "Research & Intelligence",
    "kiara":  "AI Video Generation",
    "rex":    "Remotion & Motion Graphics",
    "nova":   "Brand Visuals (Nano Banana)",
}

@app.get("/outputs", response_class=HTMLResponse)
async def outputs_page(request: Request):
    # Ensure any files from DB that are missing on disk are recreated
    db_recover_missing_files()

    active_client = request.query_params.get("client", "all")
    show_archived = active_client == "archived"

    all_jobs = db_load_all_jobs(show_archived=False)
    archived_jobs = db_load_all_jobs(show_archived=True)

    if show_archived:
        jobs = archived_jobs
    elif active_client != "all":
        jobs = [j for j in all_jobs if j.get("client") == active_client]
    else:
        jobs = all_jobs

    # ── Tab bar ──────────────────────────────────────────────────
    tabs_html = ""
    tab_defs = [
        ("all",       "All",          "#1e293b", "white"),
        ("wibiz",     "WiBiz",        "#dcfce7", "#15803d"),
        ("ai-living", "AI Living",    "#dbeafe", "#1d4ed8"),
        ("charter",   "Club Charter", "#fef9c3", "#a16207"),
        ("other",     "Other",        "#f1f5f9", "#475569"),
        ("archived",  "Archived",     "#fee2e2", "#b91c1c"),
    ]
    for key, label, bg, fg in tab_defs:
        if key == "archived":
            cnt = len(archived_jobs)
        else:
            cnt = len([j for j in all_jobs if key == "all" or j.get("client") == key])
        is_active = key == active_client
        style = (f"background:{bg};color:{fg};border:2px solid {fg}40;"
                 if not is_active else
                 f"background:#1e293b;color:white;border:2px solid #1e293b;")
        tabs_html += (
            f'<a href="/outputs?client={key}" class="tab-link" style="{style}">'
            f'{label} <span style="font-size:10px;opacity:0.7;">({cnt})</span></a> '
        )

    # ── Job cards ────────────────────────────────────────────────
    if not jobs:
        body = """<div style="text-align:center;padding:80px 0;color:#94a3b8;">
          <div style="font-size:32px;margin-bottom:12px;">📂</div>
          <div style="font-size:15px;font-weight:600;color:#475569;">No jobs yet for this client</div>
          <div style="font-size:13px;margin-top:6px;">Submit a brief on the <a href="/" style="color:#1e293b;">Brief page</a>.</div>
        </div>"""
    else:
        body = ""
        for job in jobs:
            jid        = job["id"]
            title      = job.get("title") or job.get("brief", "")[:60]
            client     = job.get("client", "other")
            pname      = job.get("project_name", "")
            brief      = job.get("brief", "")
            ts         = job.get("timestamp", job.get("created_at", ""))[:16].replace("T", " ")
            job_status = job.get("status", "complete")
            cm         = CLIENT_META.get(client, CLIENT_META["other"])
            client_label = pname if (client == "other" and pname) else cm["label"]
            is_archived  = bool(job.get("archived", 0))

            # Agents used
            s1 = job.get("stage1_outputs", {})
            s2 = job.get("cascade_outputs", {})
            all_agents = list(s1.keys()) + list(s2.keys())
            agent_chips = ""
            for ag in all_agents:
                role = AGENT_ROLES.get(ag, ag.capitalize())
                agent_chips += (
                    f'<span style="font-size:10px;font-weight:600;padding:2px 8px;border-radius:4px;'
                    f'background:#f1f5f9;color:#475569;white-space:nowrap;">'
                    f'{ag.capitalize()} — {role}</span> '
                )

            # Assembled files
            af = job.get("assembled_files", {})
            file_btns = ""
            has_files = False
            for fkey, fname in af.items():
                fm = FILE_META.get(fkey, {"label": fkey, "bg": "#f1f5f9", "fg": "#475569"})
                path = OUTPUTS_DIR / fname
                if path.exists():
                    has_files = True
                    if fname.endswith(".html"):
                        file_btns += (
                            f'<a href="/outputs/view/{fname}" target="_blank" '
                            f'style="font-size:11px;font-weight:700;padding:5px 11px;border-radius:6px;'
                            f'background:{fm["bg"]};color:{fm["fg"]};text-decoration:none;white-space:nowrap;'
                            f'border:1px solid {fm["fg"]}40;">Open {fm["label"]} →</a> '
                        )
                    file_btns += (
                        f'<a href="/outputs/download/{fname}" '
                        f'style="font-size:11px;font-weight:600;padding:5px 11px;border-radius:6px;'
                        f'background:#f8fafc;color:#64748b;text-decoration:none;white-space:nowrap;'
                        f'border:1px solid #e2e8f0;">⬇ {fname.rsplit(".",1)[-1].upper()}</a> '
                    )
            # ZIP button
            if has_files:
                file_btns += f'<a href="/outputs/zip/{jid}" class="zip-btn">⬇ ZIP</a> '

            # Archive/unarchive button
            if is_archived:
                archive_btn = f'<button onclick="unarchiveJob(\'{jid}\')" class="job-rebrief" style="color:#16a34a;">↑ Unarchive</button>'
            else:
                archive_btn = f'<button onclick="archiveJob(\'{jid}\')" class="job-rebrief" style="color:#b91c1c;">Archive</button>'

            brief_snippet = brief[:140] + "…" if len(brief) > 140 else brief
            # strip file attachment blocks from snippet
            if "\n\n--- Attached file:" in brief_snippet:
                brief_snippet = brief_snippet.split("\n\n--- Attached file:")[0] + "…"

            # Escape for HTML attributes
            title_attr = title.lower().replace('"', '')
            brief_attr = brief_snippet.lower().replace('"', '')

            _border_color = "#3b82f6" if job_status == "processing" else ("#ef4444" if job_status == "failed" else cm['fg'] + "40")
            _status_banner = ""
            if job_status == "processing":
                _status_banner = (
                    f'<div class="job-processing-banner" id="banner-{jid}">'
                    f'<span class="proc-spinner"></span>'
                    f'<span class="proc-msg" id="msg-{jid}">Generating…</span>'
                    f'</div>'
                )
            elif job_status == "failed":
                _status_banner = (
                    '<div style="margin:6px 0;padding:6px 10px;border-radius:6px;'
                    'background:#fee2e2;color:#b91c1c;font-size:12px;font-weight:600;">'
                    'Generation failed — re-brief to retry</div>'
                )

            body += f"""
<div class="job-card" id="card-{jid}" style="border-left:4px solid {_border_color};" data-id="{jid}" data-title="{title_attr}" data-brief="{brief_attr}" data-status="{job_status}">
  <div class="job-top">
    <input type="checkbox" class="job-check" onchange="onCheckChange()" title="Select for bulk action">
    <div style="flex:1;min-width:0;">
      <div style="display:flex;align-items:center;gap:8px;flex-wrap:wrap;">
        <span class="job-title">{title}</span>
        <span style="font-size:10px;font-weight:700;padding:2px 8px;border-radius:4px;background:{cm['bg']};color:{cm['fg']};white-space:nowrap;">{client_label.upper()}</span>
      </div>
      <div class="job-ts">{ts}</div>
      <div class="job-brief-text">{brief_snippet}</div>
    </div>
    <button class="star-btn" onclick="toggleStar('{jid}', this)" title="Star / pin this job">☆</button>
  </div>
  {_status_banner}
  {"<div class='job-agents'>" + agent_chips + "</div>" if agent_chips else ""}
  <div class="job-files">
    {file_btns}
    <a href="/?rebrief={jid}" class="job-rebrief">↩ Re-brief</a>
    {archive_btn}
  </div>
</div>"""

    count = len(jobs)
    total = len(all_jobs)
    return HTMLResponse(f"""<!DOCTYPE html><html lang="en"><head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>Outputs — Studio N</title>
<style>{_PAGE_STYLE}</style></head><body>
{_nav('outputs')}
<div class="container">
  <div class="page-head">
    <span class="page-title">Outputs</span>
    <span class="page-count" id="result-count">{count} job{'s' if count!=1 else ''} · {total} total</span>
  </div>
  <div class="tab-bar">{tabs_html}</div>
  <div class="search-bar">
    <input class="search-input" type="text" placeholder="Search by title or brief content…" oninput="filterJobs(this.value)">
  </div>
  <div id="jobs-list">{body}</div>
</div>
<div class="bulk-bar" id="bulk-bar">
  <span class="bulk-bar-label" id="bulk-label">0 selected</span>
  <div style="display:flex;gap:8px;">
    <button class="bulk-cancel-btn" onclick="clearSelection()">Cancel</button>
    <button class="bulk-archive-btn" onclick="bulkArchive()">Archive selected</button>
  </div>
</div>
<script>
// ── Star / Pin ────────────────────────────────────────────────
var STARRED_KEY = 'studio_starred';
function getStarred() {{
  try {{ return JSON.parse(localStorage.getItem(STARRED_KEY) || '[]'); }} catch(e) {{ return []; }}
}}
function saveStarred(arr) {{ localStorage.setItem(STARRED_KEY, JSON.stringify(arr)); }}

function toggleStar(jid, btn) {{
  var starred = getStarred();
  var idx = starred.indexOf(jid);
  if (idx === -1) {{ starred.push(jid); btn.textContent = '★'; btn.closest('.job-card').classList.add('starred'); }}
  else {{ starred.splice(idx, 1); btn.textContent = '☆'; btn.closest('.job-card').classList.remove('starred'); }}
  saveStarred(starred);
  sortCards();
}}

function sortCards() {{
  var list = document.getElementById('jobs-list');
  var cards = Array.from(list.querySelectorAll('.job-card'));
  var starred = getStarred();
  cards.sort(function(a, b) {{
    var aS = starred.includes(a.dataset.id) ? 1 : 0;
    var bS = starred.includes(b.dataset.id) ? 1 : 0;
    return bS - aS;
  }});
  cards.forEach(function(c) {{ list.appendChild(c); }});
}}

function initStars() {{
  var starred = getStarred();
  document.querySelectorAll('.job-card').forEach(function(card) {{
    var jid = card.dataset.id;
    var btn = card.querySelector('.star-btn');
    if (starred.includes(jid)) {{
      card.classList.add('starred');
      if (btn) btn.textContent = '★';
    }}
  }});
  sortCards();
}}

// ── Bulk Archive ──────────────────────────────────────────────
function onCheckChange() {{
  var checked = document.querySelectorAll('.job-check:checked');
  var bar = document.getElementById('bulk-bar');
  var label = document.getElementById('bulk-label');
  if (checked.length > 0) {{
    bar.classList.add('visible');
    label.textContent = checked.length + ' selected';
  }} else {{
    bar.classList.remove('visible');
  }}
}}

function clearSelection() {{
  document.querySelectorAll('.job-check:checked').forEach(function(cb) {{ cb.checked = false; }});
  document.getElementById('bulk-bar').classList.remove('visible');
}}

async function archiveJob(jid) {{
  await fetch('/api/job/' + jid + '/archive', {{method:'POST'}});
  var card = document.getElementById('card-' + jid);
  if (card) card.remove();
  updateCount();
}}

async function unarchiveJob(jid) {{
  await fetch('/api/job/' + jid + '/unarchive', {{method:'POST'}});
  var card = document.getElementById('card-' + jid);
  if (card) card.remove();
  updateCount();
}}

async function bulkArchive() {{
  var checked = document.querySelectorAll('.job-check:checked');
  var ids = [];
  checked.forEach(function(cb) {{ ids.push(cb.closest('.job-card').dataset.id); }});
  for (var i = 0; i < ids.length; i++) {{
    await fetch('/api/job/' + ids[i] + '/archive', {{method:'POST'}});
    var card = document.getElementById('card-' + ids[i]);
    if (card) card.remove();
  }}
  document.getElementById('bulk-bar').classList.remove('visible');
  updateCount();
}}

function updateCount() {{
  var visible = document.querySelectorAll('.job-card:not([style*="display: none"])').length;
  var el = document.getElementById('result-count');
  if (el) el.textContent = visible + ' job' + (visible !== 1 ? 's' : '') + ' · {total} total';
}}

// ── Search ────────────────────────────────────────────────────
function filterJobs(q) {{
  q = q.toLowerCase().trim();
  var cards = document.querySelectorAll('.job-card');
  var shown = 0;
  cards.forEach(function(c) {{
    var match = !q || c.dataset.title.includes(q) || c.dataset.brief.includes(q);
    c.style.display = match ? '' : 'none';
    if (match) shown++;
  }});
  var el = document.getElementById('result-count');
  if (el) el.textContent = shown + ' job' + (shown !== 1 ? 's' : '') + ' · {total} total';
}}

// Init on load
initStars();

// ── In-progress job polling ───────────────────────────────────
(function() {{
  var processingCards = document.querySelectorAll('.job-card[data-status="processing"]');
  if (processingCards.length === 0) return;

  var pollers = {{}};

  function pollJob(jid) {{
    fetch('/api/job/' + jid + '/progress')
      .then(function(r) {{ return r.json(); }})
      .then(function(data) {{
        var msgEl = document.getElementById('msg-' + jid);
        var card  = document.getElementById('card-' + jid);
        if (data.status === 'processing') {{
          if (msgEl && data.message) msgEl.textContent = data.message;
        }} else if (data.status === 'complete' || data.status === 'failed') {{
          clearInterval(pollers[jid]);
          delete pollers[jid];
          // Reload the page so the card renders with full outputs
          window.location.reload();
        }}
      }})
      .catch(function() {{/* network blip — keep polling */}});
  }}

  processingCards.forEach(function(card) {{
    var jid = card.dataset.id;
    pollers[jid] = setInterval(function() {{ pollJob(jid); }}, 4000);
    pollJob(jid); // immediate first check
  }});
}})();
</script>
<style>
.job-processing-banner {{
  display:flex;align-items:center;gap:8px;
  margin:6px 0;padding:8px 12px;
  border-radius:7px;background:#eff6ff;
  border:1px solid #bfdbfe;
  font-size:12px;font-weight:600;color:#1d4ed8;
}}
.proc-spinner {{
  display:inline-block;width:12px;height:12px;
  border:2px solid #93c5fd;border-top-color:#1d4ed8;
  border-radius:50%;animation:proc-spin 0.8s linear infinite;flex-shrink:0;
}}
@keyframes proc-spin {{ to {{ transform:rotate(360deg); }} }}
.proc-msg {{ flex:1; }}
</style>
</body></html>""")

# ── Legacy file-based outputs (kept for backward compat) ──────────

def _parse_output_file(path: Path) -> dict:
    info = {"filename": path.name, "size_kb": round(path.stat().st_size/1024,1),
            "brief": path.stem, "date": "", "is_html": path.suffix == ".html",
            "is_json": path.suffix == ".json", "label": ""}
    for l in ["website", "one-pager", "social-pack", "video-brief", "canva"]:
        if l in path.name: info["label"] = l; break
    try:
        text = path.read_text(encoding="utf-8")
        if path.suffix == ".md":
            if m := re.search(r"\*\*Brief:\*\* (.+)", text): info["brief"] = m.group(1).strip()
            if d := re.search(r"\*\*Date:\*\* (.+)",  text): info["date"]  = d.group(1).strip()
        elif path.suffix == ".html":
            if t := re.search(r"<title>(.+?)</title>", text): info["brief"] = t.group(1).strip()
        elif path.suffix == ".json":
            try:
                data = json.loads(text)
                info["brief"] = data.get("brand", {}).get("name", path.stem)
            except Exception:
                pass
    except Exception:
        pass
    return info

@app.get("/outputs", response_class=HTMLResponse)
async def outputs_page():
    files = [_parse_output_file(f) for f in
             sorted(OUTPUTS_DIR.glob("*"), key=lambda f: f.stat().st_mtime, reverse=True)
             if f.suffix in (".md", ".html", ".json")]
    rows = ""
    for f in files:
        label_styles = {
            "website":    "#dbeafe:#1d4ed8",
            "one-pager":  "#d1fae5:#065f46",
            "canva":      "#ede9fe:#6d28d9",
            "social-pack":"#fef3c7:#92400e",
        }
        lbg, lfg = label_styles.get(f["label"], "#f1f5f9:#475569").split(":") if f["label"] else ("#f1f5f9", "#475569")
        lb = (f'<span style="font-size:10px;font-weight:700;padding:2px 7px;border-radius:4px;'
              f'background:{lbg};color:{lfg};margin-left:8px;">{f["label"].upper()}</span>'
              if f["label"] else "")
        if f["is_html"]:
            view = f'<a class="out-btn out-btn-view" href="/outputs/view/{f["filename"]}" target="_blank">View</a>'
        elif f["is_json"]:
            view = f'<a class="out-btn out-btn-json" href="/outputs/download/{f["filename"]}">⬇ JSON</a>'
        else:
            view = f'<a class="out-btn out-btn-pdf" href="/outputs/view/{f["filename"]}" target="_blank">View</a>'
        ext = f["filename"].rsplit(".",1)[-1].upper()
        rows += f"""<div class="output-row">
          <div class="output-meta">
            <span class="output-date">{f['date'] or '—'}</span>
            <span class="output-size">{f['size_kb']} KB</span>
          </div>
          <div class="output-brief">{f['brief']}{lb}</div>
          <div class="output-actions">
            <a class="out-btn" href="/outputs/download/{f['filename']}">⬇ {ext}</a>
            {view}
          </div>
        </div>"""
    if not files:
        rows = """<div style="text-align:center;padding:60px 0;color:#94a3b8;">
          <div style="font-size:32px;margin-bottom:12px;">📂</div>
          <div style="font-size:15px;font-weight:600;color:#475569;">No outputs yet</div>
          <div style="font-size:13px;margin-top:6px;">Submit a brief on the <a href="/" style="color:#1e293b;">Brief page</a>.</div>
        </div>"""
    count = len(files)
    return HTMLResponse(f"""<!DOCTYPE html><html lang="en"><head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>Outputs — Studio N</title>
<style>
  {_PAGE_STYLE}
  .page-head {{display:flex;align-items:baseline;justify-content:space-between;margin-bottom:24px;}}
  .page-title {{font-size:20px;font-weight:700;}} .page-count {{font-size:13px;color:#94a3b8;}}
  .output-row {{background:white;border-radius:10px;padding:18px 20px;margin-bottom:10px;
    box-shadow:0 1px 3px rgba(0,0,0,0.07);display:flex;align-items:center;gap:16px;transition:box-shadow 0.15s;}}
  .output-row:hover {{box-shadow:0 2px 8px rgba(0,0,0,0.1);}}
  .output-meta {{display:flex;flex-direction:column;gap:3px;min-width:130px;}}
  .output-date {{font-size:12px;font-weight:600;color:#475569;}} .output-size {{font-size:11px;color:#94a3b8;}}
  .output-brief {{flex:1;font-size:13px;color:#334155;line-height:1.5;}}
  .output-actions {{display:flex;gap:6px;flex-shrink:0;}}
  .out-btn {{font-size:11px;font-weight:600;padding:5px 11px;border-radius:6px;border:1px solid #e2e8f0;
    color:#475569;text-decoration:none;transition:all 0.15s;white-space:nowrap;}}
  .out-btn:hover {{background:#f8fafc;color:#1e293b;}}
  .out-btn-pdf {{border-color:#bfdbfe;color:#2563eb;}} .out-btn-pdf:hover {{background:#eff6ff;}}
  .out-btn-view {{border-color:#bbf7d0;color:#16a34a;font-weight:700;}} .out-btn-view:hover {{background:#f0fdf4;}}
  .out-btn-json {{border-color:#e9d5ff;color:#7c3aed;font-weight:700;}} .out-btn-json:hover {{background:#faf5ff;}}
</style></head><body>
{_nav('outputs')}
<div class="container">
  <div class="page-head">
    <span class="page-title">Outputs</span>
    <span class="page-count">{count} file{'s' if count!=1 else ''}</span>
  </div>
  {rows}
</div></body></html>""")

@app.get("/outputs/download/{filename}")
async def download_output(filename: str):
    safe = Path(filename).name
    path = OUTPUTS_DIR / safe
    if not path.exists() or path.suffix not in (".md", ".html", ".json"):
        return JSONResponse({"error": "File not found"}, status_code=404)
    mime = {"html": "text/html", "md": "text/markdown", "json": "application/json"}.get(
        path.suffix.lstrip("."), "application/octet-stream")
    return StreamingResponse(iter([path.read_bytes()]), media_type=mime,
                             headers={"Content-Disposition": f'attachment; filename="{safe}"'})

@app.get("/outputs/view/{filename}", response_class=HTMLResponse)
async def view_output(filename: str):
    safe = Path(filename).name
    path = OUTPUTS_DIR / safe
    if not path.exists() or path.suffix not in (".md", ".html", ".json"):
        return HTMLResponse("<h1>File not found</h1>", status_code=404)
    content = path.read_text(encoding="utf-8")
    if path.suffix == ".html":
        return HTMLResponse(content)
    bm = re.search(r"\*\*Brief:\*\* (.+)", content)
    dm = re.search(r"\*\*Date:\*\* (.+)",  content)
    return HTMLResponse(_print_page(
        safe,
        bm.group(1).strip() if bm else safe,
        dm.group(1).strip() if dm else "",
        _md_to_html(content)))

@app.get("/health")
async def health():
    return {"status": "ok"}


# ── Campaign routes ───────────────────────────────────────────────

@app.post("/api/campaign")
async def campaign_start(request: Request):
    body          = await request.json()
    campaign_id   = str(uuid.uuid4())
    client        = body.get("client", "wibiz")
    platform      = body.get("platform", "instagram")
    duration      = int(body.get("duration_days", 30))
    style         = body.get("style_brief", "")
    accounts      = body.get("reference_accounts", [])
    urls          = body.get("reference_urls", [])
    mix           = body.get("content_mix", {"video": 0.3, "static": 0.5, "carousel": 0.2})

    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        INSERT INTO campaigns (id, client, platform, duration_days, style_brief, reference_accounts, status)
        VALUES (?,?,?,?,?,?,'processing')
    """, (campaign_id, client, platform, duration, style, json.dumps(accounts)))
    conn.commit()
    conn.close()

    _campaign_status[campaign_id] = {"step": "starting", "message": "Campaign initialising..."}
    asyncio.create_task(run_campaign(
        campaign_id, client, platform, duration, style, accounts, urls, mix
    ))
    return {"campaign_id": campaign_id}

@app.get("/api/campaign/{campaign_id}/status")
async def get_campaign_status(campaign_id: str):
    cached = _campaign_status.get(campaign_id)
    if cached:
        return cached
    conn = sqlite3.connect(DB_PATH)
    row  = conn.execute(
        "SELECT status, current_step, calendar_json, video_job_ids FROM campaigns WHERE id=?",
        (campaign_id,)
    ).fetchone()
    conn.close()
    if not row:
        return JSONResponse({"error": "not found"}, status_code=404)
    cal = {}
    try: cal = json.loads(row[2] or "{}")
    except Exception: pass
    return {"step": row[0] or "unknown", "message": row[1] or "", "calendar": cal}

@app.get("/api/campaigns")
async def list_campaigns():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    rows = conn.execute(
        "SELECT id, client, platform, duration_days, style_brief, status, current_step, created_at "
        "FROM campaigns ORDER BY created_at DESC LIMIT 50"
    ).fetchall()
    conn.close()
    return {"campaigns": [dict(r) for r in rows]}

@app.get("/campaign", response_class=HTMLResponse)
async def campaign_page(request: Request):
    if not request.session.get("authenticated"):
        return RedirectResponse("/login", status_code=302)
    client_options = "".join(
        f'<option value="{k}"{" selected" if k=="wibiz" else ""}>{v["label"]}</option>'
        for k, v in CLIENT_META.items() if k != "other"
    )
    return HTMLResponse(f"""<!DOCTYPE html><html lang="en"><head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>Campaign Planner — Studio N v7.0</title>
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap" rel="stylesheet">
<style>
{_PAGE_STYLE}
.camp-wrap {{max-width:820px;margin:0 auto;padding:32px 24px 80px;}}
.camp-form {{background:white;border-radius:14px;padding:28px 32px;box-shadow:0 1px 4px rgba(0,0,0,0.08);margin-bottom:24px;}}
.camp-title {{font-size:18px;font-weight:800;color:#1e293b;margin-bottom:20px;}}
.field-row {{display:grid;grid-template-columns:1fr 1fr;gap:16px;margin-bottom:16px;}}
.field-row-3 {{display:grid;grid-template-columns:1fr 1fr 1fr;gap:16px;margin-bottom:16px;}}
.field {{display:flex;flex-direction:column;gap:5px;margin-bottom:16px;}}
.field-label {{font-size:12px;font-weight:600;color:#475569;text-transform:uppercase;letter-spacing:0.04em;}}
.field-input, .field-select {{border:1px solid #e2e8f0;border-radius:7px;padding:9px 12px;font-size:14px;color:#1e293b;background:white;outline:none;transition:border 0.15s;width:100%;}}
.field-input:focus, .field-select:focus {{border-color:#6366f1;}}
.field-textarea {{border:1px solid #e2e8f0;border-radius:7px;padding:9px 12px;font-size:13px;color:#1e293b;background:white;outline:none;width:100%;resize:vertical;min-height:80px;font-family:inherit;}}
.field-textarea:focus {{border-color:#6366f1;}}
.platform-grid {{display:flex;gap:8px;flex-wrap:wrap;margin-bottom:16px;}}
.plat-btn {{padding:7px 14px;border-radius:7px;border:1px solid #e2e8f0;background:white;font-size:12px;font-weight:600;cursor:pointer;transition:all 0.15s;color:#475569;}}
.plat-btn.active {{border-color:#6366f1;background:#eff6ff;color:#4f46e5;}}
.mix-row {{display:flex;gap:12px;align-items:center;margin-bottom:16px;}}
.mix-chip {{display:flex;align-items:center;gap:6px;font-size:12px;color:#475569;}}
.mix-chip input {{width:48px;padding:4px 7px;border:1px solid #e2e8f0;border-radius:5px;font-size:12px;}}
.btn-launch {{background:linear-gradient(135deg,#6366f1,#8b5cf6);color:white;border:none;padding:12px 28px;border-radius:8px;font-size:14px;font-weight:700;cursor:pointer;transition:all 0.15s;}}
.btn-launch:hover {{opacity:0.9;transform:translateY(-1px);}}
.btn-launch:disabled {{opacity:0.5;cursor:not-allowed;transform:none;}}
.progress-box {{background:white;border-radius:14px;padding:24px;box-shadow:0 1px 4px rgba(0,0,0,0.08);margin-bottom:24px;display:none;}}
.progress-title {{font-size:14px;font-weight:700;color:#1e293b;margin-bottom:12px;}}
.step-list {{display:flex;flex-direction:column;gap:6px;}}
.step {{display:flex;align-items:center;gap:8px;font-size:13px;color:#64748b;padding:6px 0;}}
.step.active {{color:#4f46e5;font-weight:600;}}
.step.done {{color:#16a34a;}}
.step.failed {{color:#dc2626;}}
.step-icon {{width:18px;height:18px;border-radius:50%;display:flex;align-items:center;justify-content:center;font-size:10px;flex-shrink:0;}}
.step-icon.active {{background:#eff6ff;border:1px solid #6366f1;}}
.step-icon.done {{background:#dcfce7;color:#16a34a;}}
.step-icon.failed {{background:#fee2e2;color:#dc2626;}}
.step-icon.pending {{background:#f8fafc;border:1px solid #e2e8f0;}}
.spin {{display:inline-block;width:10px;height:10px;border:2px solid #c7d2fe;border-top-color:#4f46e5;border-radius:50%;animation:spin 0.8s linear infinite;}}
@keyframes spin {{to{{transform:rotate(360deg)}}}}
.calendar-wrap {{background:white;border-radius:14px;padding:24px;box-shadow:0 1px 4px rgba(0,0,0,0.08);margin-bottom:24px;display:none;}}
.cal-header {{display:flex;align-items:center;justify-content:space-between;margin-bottom:18px;}}
.cal-title {{font-size:16px;font-weight:800;color:#1e293b;}}
.cal-meta {{font-size:12px;color:#94a3b8;}}
.cal-actions {{display:flex;gap:8px;}}
.btn-sm {{font-size:11px;font-weight:600;padding:5px 12px;border-radius:6px;border:1px solid #e2e8f0;background:white;cursor:pointer;color:#475569;}}
.btn-sm:hover {{background:#f8fafc;}}
.cal-table {{width:100%;border-collapse:collapse;font-size:12px;}}
.cal-table th {{padding:8px 10px;text-align:left;font-size:11px;font-weight:700;color:#64748b;text-transform:uppercase;letter-spacing:0.04em;border-bottom:2px solid #f1f5f9;}}
.cal-table td {{padding:9px 10px;border-bottom:1px solid #f1f5f9;vertical-align:top;}}
.cal-table tr:hover td {{background:#fafbff;}}
.type-badge {{display:inline-block;padding:2px 7px;border-radius:4px;font-size:10px;font-weight:700;text-transform:uppercase;}}
.type-video   {{background:#dbeafe;color:#1d4ed8;}}
.type-static  {{background:#dcfce7;color:#16a34a;}}
.type-carousel {{background:#fef3c7;color:#b45309;}}
.video-jobs-wrap {{background:white;border-radius:14px;padding:24px;box-shadow:0 1px 4px rgba(0,0,0,0.08);display:none;}}
.vj-title {{font-size:16px;font-weight:800;color:#1e293b;margin-bottom:14px;}}
.vj-list {{display:grid;grid-template-columns:repeat(auto-fill,minmax(200px,1fr));gap:10px;}}
.vj-card {{border:1px solid #e2e8f0;border-radius:8px;padding:10px;font-size:12px;}}
.vj-id {{font-family:monospace;color:#6366f1;font-size:10px;margin-bottom:4px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;}}
.past-list {{display:flex;flex-direction:column;gap:8px;}}
.past-card {{border:1px solid #e2e8f0;border-radius:8px;padding:12px 16px;display:flex;align-items:center;gap:12px;cursor:pointer;transition:all 0.15s;}}
.past-card:hover {{border-color:#6366f1;background:#fafbff;}}
</style>
</head><body>
{_nav('campaign')}
<div class="camp-wrap">
  <div class="camp-form">
    <div class="camp-title">Campaign Planner</div>
    <div class="field-row">
      <div class="field">
        <label class="field-label">Client</label>
        <select class="field-select" id="camp-client">{client_options}</select>
      </div>
      <div class="field">
        <label class="field-label">Duration (days)</label>
        <select class="field-select" id="camp-duration">
          <option value="3">3 days</option>
          <option value="7">7 days</option>
          <option value="14">14 days</option>
          <option value="30" selected>30 days</option>
          <option value="60">60 days</option>
          <option value="90">90 days</option>
        </select>
      </div>
    </div>
    <div class="field">
      <label class="field-label">Platform</label>
      <div class="platform-grid" id="platform-grid">
        <button class="plat-btn active" onclick="setPlatform('instagram',this)">Instagram</button>
        <button class="plat-btn" onclick="setPlatform('tiktok',this)">TikTok</button>
        <button class="plat-btn" onclick="setPlatform('youtube',this)">YouTube</button>
        <button class="plat-btn" onclick="setPlatform('facebook',this)">Facebook</button>
        <button class="plat-btn" onclick="setPlatform('linkedin',this)">LinkedIn</button>
      </div>
    </div>
    <div class="field">
      <label class="field-label">Style Brief</label>
      <textarea class="field-textarea" id="camp-style" placeholder="Describe the visual style, tone, and feel. E.g. 'Bold and fast-paced like Gymshark, educational hooks, heavy use of text overlays, trending audio'"></textarea>
    </div>
    <div class="field">
      <label class="field-label">Reference Accounts (one per line, with or without @)</label>
      <textarea class="field-textarea" id="camp-accounts" style="min-height:64px;" placeholder="@competitor1&#10;@inspiration_account&#10;accountname"></textarea>
    </div>
    <div class="field">
      <label class="field-label">Content Mix</label>
      <div class="mix-row">
        <div class="mix-chip">Video <input type="number" id="mix-video" value="30" min="0" max="100">%</div>
        <div class="mix-chip">Static <input type="number" id="mix-static" value="50" min="0" max="100">%</div>
        <div class="mix-chip">Carousel <input type="number" id="mix-carousel" value="20" min="0" max="100">%</div>
      </div>
    </div>
    <div style="display:flex;align-items:center;gap:14px;">
      <button class="btn-launch" onclick="launchCampaign()" id="launch-btn">
        Launch Campaign
      </button>
      <span id="launch-hint" style="font-size:12px;color:#94a3b8;">Zara researches → Marcus plans → Priya + Dante write → Kiara fires to Arcads</span>
    </div>
  </div>

  <div class="progress-box" id="progress-box">
    <div class="progress-title" id="progress-title">Running campaign...</div>
    <div class="step-list" id="step-list">
      <div class="step" id="st-research"><div class="step-icon pending" id="si-research">·</div><span>Zara — research &amp; competitor scrape</span></div>
      <div class="step" id="st-planning"><div class="step-icon pending" id="si-planning">·</div><span>Marcus — build content calendar</span></div>
      <div class="step" id="st-writing"><div class="step-icon pending" id="si-writing">·</div><span>Priya + Dante — captions &amp; scripts (parallel)</span></div>
      <div class="step" id="st-video_specs"><div class="step-icon pending" id="si-video_specs">·</div><span>Kiara — Arcads video specs</span></div>
      <div class="step" id="st-generating_videos"><div class="step-icon pending" id="si-generating_videos">·</div><span>Auto-fire video jobs to Arcads</span></div>
      <div class="step" id="st-assembling"><div class="step-icon pending" id="si-assembling">·</div><span>Assembling final calendar</span></div>
    </div>
    <div id="progress-msg" style="margin-top:10px;font-size:12px;color:#6366f1;"></div>
  </div>

  <div class="calendar-wrap" id="calendar-wrap">
    <div class="cal-header">
      <div>
        <div class="cal-title" id="cal-title">Content Calendar</div>
        <div class="cal-meta" id="cal-meta"></div>
      </div>
      <div class="cal-actions">
        <button class="btn-sm" onclick="downloadCalendar()">Download JSON</button>
        <button class="btn-sm" onclick="copyCalendar()">Copy CSV</button>
      </div>
    </div>
    <table class="cal-table">
      <thead><tr>
        <th>Day</th><th>Type</th><th>Theme</th><th>Hook</th><th>Caption</th><th>CTA</th>
      </tr></thead>
      <tbody id="cal-body"></tbody>
    </table>
  </div>

  <div class="video-jobs-wrap" id="vj-wrap">
    <div class="vj-title">Video Jobs Queued in Arcads</div>
    <div class="vj-list" id="vj-list"></div>
    <p style="margin-top:10px;font-size:12px;color:#94a3b8;">These are processing in Arcads now. Check Video Studio for results.</p>
  </div>

  <div class="camp-form" style="margin-top:24px;">
    <div class="camp-title" style="font-size:15px;">Past Campaigns</div>
    <div class="past-list" id="past-list"><div style="font-size:13px;color:#94a3b8;">Loading...</div></div>
  </div>
</div>

<script>
var _platform = 'instagram';
var _activeCampaignId = null;
var _poller = null;
var _calendarData = null;

var STEP_ORDER = ['research','planning','writing','video_specs','generating_videos','assembling','complete'];

function setPlatform(p, btn) {{
  _platform = p;
  document.querySelectorAll('.plat-btn').forEach(function(b) {{ b.classList.remove('active'); }});
  btn.classList.add('active');
}}

function launchCampaign() {{
  var style = document.getElementById('camp-style').value.trim();
  if (!style) {{ alert('Add a style brief first.'); return; }}

  var accountsRaw = document.getElementById('camp-accounts').value.trim();
  var accounts = accountsRaw ? accountsRaw.split('\\n').map(function(a) {{ return a.trim().replace(/^@/,''); }}).filter(Boolean) : [];

  var mv = parseInt(document.getElementById('mix-video').value) || 30;
  var ms = parseInt(document.getElementById('mix-static').value) || 50;
  var mc = parseInt(document.getElementById('mix-carousel').value) || 20;
  var total = mv + ms + mc;

  var payload = {{
    client: document.getElementById('camp-client').value,
    platform: _platform,
    duration_days: parseInt(document.getElementById('camp-duration').value),
    style_brief: style,
    reference_accounts: accounts,
    reference_urls: [],
    content_mix: {{
      video: mv / total,
      static: ms / total,
      carousel: mc / total,
    }},
  }};

  document.getElementById('launch-btn').disabled = true;
  document.getElementById('launch-hint').textContent = 'Campaign running — do not navigate away...';
  document.getElementById('progress-box').style.display = '';
  document.getElementById('calendar-wrap').style.display = 'none';
  document.getElementById('vj-wrap').style.display = 'none';
  resetSteps();

  fetch('/api/campaign', {{
    method: 'POST',
    headers: {{'Content-Type':'application/json'}},
    body: JSON.stringify(payload),
  }})
  .then(function(r) {{ return r.json(); }})
  .then(function(data) {{
    _activeCampaignId = data.campaign_id;
    _poller = setInterval(function() {{ pollCampaign(_activeCampaignId); }}, 5000);
    pollCampaign(_activeCampaignId);
  }})
  .catch(function(e) {{
    alert('Failed to start campaign: ' + e);
    document.getElementById('launch-btn').disabled = false;
  }});
}}

function resetSteps() {{
  STEP_ORDER.forEach(function(s) {{
    var si = document.getElementById('si-' + s);
    var st = document.getElementById('st-' + s);
    if (si) {{ si.className = 'step-icon pending'; si.textContent = '·'; }}
    if (st) st.className = 'step';
  }});
}}

function markStep(current) {{
  var idx = STEP_ORDER.indexOf(current);
  STEP_ORDER.forEach(function(s, i) {{
    var si = document.getElementById('si-' + s);
    var st = document.getElementById('st-' + s);
    if (!si) return;
    if (i < idx) {{
      si.className = 'step-icon done'; si.textContent = '✓';
      if (st) st.className = 'step done';
    }} else if (i === idx) {{
      si.className = 'step-icon active'; si.innerHTML = '<span class="spin"></span>';
      if (st) st.className = 'step active';
    }} else {{
      si.className = 'step-icon pending'; si.textContent = '·';
      if (st) st.className = 'step';
    }}
  }});
}}

function pollCampaign(cid) {{
  fetch('/api/campaign/' + cid + '/status')
    .then(function(r) {{ return r.json(); }})
    .then(function(data) {{
      var step = data.step || '';
      var msg  = data.message || '';
      document.getElementById('progress-msg').textContent = msg;
      document.getElementById('progress-title').textContent =
        step === 'complete' ? 'Campaign complete!' :
        step === 'failed'   ? 'Campaign failed' : 'Running campaign...';

      if (step === 'failed') {{
        markStep('failed');
        clearInterval(_poller);
        document.getElementById('launch-btn').disabled = false;
        document.getElementById('launch-hint').textContent = 'Campaign failed — check the message above.';
        return;
      }}

      markStep(step);

      if (step === 'complete' && data.calendar) {{
        clearInterval(_poller);
        document.getElementById('launch-btn').disabled = false;
        document.getElementById('launch-hint').textContent = 'Campaign complete.';
        // Mark all steps done
        STEP_ORDER.forEach(function(s) {{
          var si = document.getElementById('si-' + s);
          var st = document.getElementById('st-' + s);
          if (si) {{ si.className = 'step-icon done'; si.textContent = '✓'; }}
          if (st) st.className = 'step done';
        }});
        renderCalendar(data.calendar);
        renderVideoJobs(data.calendar.video_job_ids || []);
        loadPastCampaigns();
      }}
    }})
    .catch(function() {{}});
}}

function renderCalendar(cal) {{
  _calendarData = cal;
  var days = cal.days || [];
  document.getElementById('cal-title').textContent =
    (cal.platform || 'Campaign').charAt(0).toUpperCase() + (cal.platform || '').slice(1) + ' Content Calendar';
  document.getElementById('cal-meta').textContent =
    days.length + ' days · ' + (cal.client || '') + ' · ' + (cal.style || '');

  var rows = days.map(function(d) {{
    var typeCls = 'type-' + (d.content_type || 'static');
    var caption = (d.caption || d.key_message || '').replace(/</g,'&lt;').substring(0,120);
    return '<tr>' +
      '<td><strong>' + d.day + '</strong></td>' +
      '<td><span class="type-badge ' + typeCls + '">' + (d.content_type || '') + '</span></td>' +
      '<td>' + (d.theme || '').replace(/</g,'&lt;') + '</td>' +
      '<td style="max-width:160px;">' + (d.hook || '').replace(/</g,'&lt;').substring(0,80) + '</td>' +
      '<td style="max-width:200px;font-size:11px;">' + caption + '</td>' +
      '<td style="font-size:11px;">' + (d.cta || '').replace(/</g,'&lt;').substring(0,60) + '</td>' +
      '</tr>';
  }}).join('');
  document.getElementById('cal-body').innerHTML = rows || '<tr><td colspan="6" style="text-align:center;color:#94a3b8;padding:20px;">No calendar data</td></tr>';
  document.getElementById('calendar-wrap').style.display = '';
}}

function renderVideoJobs(ids) {{
  if (!ids || !ids.length) return;
  var html = ids.map(function(id) {{
    return '<div class="vj-card">' +
      '<div class="vj-id">' + id + '</div>' +
      '<div style="font-size:11px;color:#64748b;">Status: pending</div>' +
      '<a href="/video-studio" style="font-size:11px;color:#6366f1;">View in Video Studio</a>' +
      '</div>';
  }}).join('');
  document.getElementById('vj-list').innerHTML = html;
  document.getElementById('vj-wrap').style.display = '';
}}

function downloadCalendar() {{
  if (!_calendarData) return;
  var blob = new Blob([JSON.stringify(_calendarData, null, 2)], {{type:'application/json'}});
  var a = document.createElement('a'); a.href = URL.createObjectURL(blob);
  a.download = 'campaign-' + (_calendarData.platform || 'calendar') + '.json'; a.click();
}}

function copyCalendar() {{
  if (!_calendarData || !_calendarData.days) return;
  var rows = [['Day','Type','Theme','Hook','Caption','Hashtags','CTA']];
  _calendarData.days.forEach(function(d) {{
    rows.push([
      d.day, d.content_type || '', d.theme || '', d.hook || '',
      (d.caption || '').replace(/\\n/g,' '),
      (d.hashtags || []).join(' '),
      d.cta || ''
    ]);
  }});
  var csv = rows.map(function(r) {{ return r.map(function(c) {{ return '"' + String(c).replace(/"/g,'""') + '"'; }}).join(','); }}).join('\\n');
  navigator.clipboard.writeText(csv).then(function() {{ alert('CSV copied to clipboard!'); }});
}}

function loadPastCampaigns() {{
  fetch('/api/campaigns')
    .then(function(r) {{ return r.json(); }})
    .then(function(data) {{
      var camps = data.campaigns || [];
      if (!camps.length) {{
        document.getElementById('past-list').innerHTML = '<div style="font-size:13px;color:#94a3b8;">No past campaigns yet.</div>';
        return;
      }}
      var html = camps.map(function(c) {{
        var statusColor = c.status === 'complete' ? '#16a34a' : c.status === 'failed' ? '#dc2626' : '#6366f1';
        return '<div class="past-card" onclick="loadPastCampaign(\\'' + c.id + '\\'')">' +
          '<div style="flex:1;">' +
          '<div style="font-size:13px;font-weight:600;color:#1e293b;">' + (c.platform || '') + ' · ' + (c.duration_days || '') + ' days</div>' +
          '<div style="font-size:12px;color:#64748b;">' + (c.style_brief || '').substring(0,80) + '</div>' +
          '<div style="font-size:11px;color:#94a3b8;">' + (c.created_at || '').substring(0,16) + '</div>' +
          '</div>' +
          '<span style="font-size:11px;font-weight:700;color:' + statusColor + ';">' + c.status + '</span>' +
          '</div>';
      }}).join('');
      document.getElementById('past-list').innerHTML = html;
    }})
    .catch(function() {{}});
}}

function loadPastCampaign(cid) {{
  fetch('/api/campaign/' + cid + '/status')
    .then(function(r) {{ return r.json(); }})
    .then(function(data) {{
      if (data.calendar && data.calendar.days) {{
        document.getElementById('progress-box').style.display = 'none';
        renderCalendar(data.calendar);
        renderVideoJobs(data.calendar.video_job_ids || []);
      }}
    }});
}}

// Load past campaigns on page load
loadPastCampaigns();
</script>
</body></html>""")


# ── Video Studio ──────────────────────────────────────────────────

@app.get("/video-studio", response_class=HTMLResponse)
async def video_studio_page(request: Request):
    try:
        return templates.TemplateResponse(request=request, name="video_studio.html")
    except Exception as e:
        logging.error("video_studio render error: %s", traceback.format_exc())
        return HTMLResponse(f"<pre>Error: {e}</pre>", status_code=500)

@app.get("/api/video-studio/debug-auth")
async def vs_debug_auth():
    cid = ARCADS_CLIENT_ID
    sec = ARCADS_CLIENT_SECRET
    masked = sec[:3] + "..." + sec[-3:] if len(sec) > 6 else "TOO_SHORT"
    import base64
    creds = base64.b64encode(f"{cid}:{sec}".encode()).decode()
    async with httpx.AsyncClient() as client:
        r = await client.get(f"{ARCADS_BASE_URL}/v1/products",
                             headers={"Authorization": f"Basic {creds}"}, timeout=15)
    return {
        "client_id_len": len(cid),
        "secret_len": len(sec),
        "secret_masked": masked,
        "status": r.status_code,
        "response": r.text[:200],
    }

@app.get("/api/video-studio/debug-upload")
async def vs_debug_upload():
    """Test the full presigned upload pipeline with a tiny dummy file."""
    try:
        dummy = b"test-upload-" + str(uuid.uuid4()).encode()
        import base64
        creds = base64.b64encode(f"{ARCADS_CLIENT_ID}:{ARCADS_CLIENT_SECRET}".encode()).decode()
        async with httpx.AsyncClient() as client:
            # Step 1: get presigned URL
            r1 = await client.post(f"{ARCADS_BASE_URL}/v1/file-upload/get-presigned-url",
                                   json={"fileType": "video/mp4"},
                                   headers={"Authorization": f"Basic {creds}", "Content-Type": "application/json"},
                                   timeout=15)
            if not r1.is_success:
                return {"step": "presigned", "status": r1.status_code, "body": r1.text[:400]}
            data = r1.json()
            upload_url = data.get("presignedUrl") or ""
            file_path  = data.get("filePath") or ""
            # Step 2: PUT dummy file
            r2 = await client.put(upload_url, content=dummy,
                                  headers={"Content-Type": "video/mp4"}, timeout=30)
            return {
                "step": "complete",
                "presigned_status": r1.status_code,
                "s3_put_status": r2.status_code,
                "filePath": file_path,
                "presigned_response_keys": list(data.keys()),
            }
    except Exception as e:
        return {"error": str(e)}

@app.get("/api/video-studio/products")
async def vs_get_products():
    try:
        products = await arcads_get_products()
        return {"products": products}
    except Exception as e:
        return JSONResponse({"error": str(e), "products": []}, status_code=200)

@app.post("/api/video-studio/generate")
async def vs_generate(request: Request):
    try:
        body = await request.json()
        client_val = body.get("client", "other")
        model      = body.get("model", "seedance-2.0")
        prompt     = body.get("prompt", "")
        product_id = body.get("productId", "")
        formats    = body.get("formats", ["9:16"])
        variations = int(body.get("variations", 1))
        duration   = body.get("duration")

        if not ARCADS_CLIENT_ID:
            return JSONResponse({"error": "ARCADS_CLIENT_ID not configured"}, status_code=400)
        if not prompt:
            return JSONResponse({"error": "Prompt is required"}, status_code=400)

        # Auto-resolve product ID if not provided
        if not product_id:
            try:
                products = await arcads_get_products()
                logging.info("Auto-resolve products: %s", products)
                if products:
                    product_id = products[0].get("id", "")
            except Exception as pe:
                logging.error("Auto-resolve product error: %s", pe)

        logging.info("Using productId: %s", product_id)
        if not product_id:
            return JSONResponse({"error": "No product ID — add a product in your Arcads account first"}, status_code=400)

        job_id = str(uuid.uuid4())
        jobs_created = []

        for fmt in formats:
            for _ in range(variations):
                vid_job_id = str(uuid.uuid4())
                payload = {
                    "model": model,
                    "productId": product_id,
                    "prompt": prompt,
                }
                if fmt and fmt != "auto":
                    if model == "seedance-2.0":
                        payload["resolution"] = SEEDANCE2_RESOLUTION_MAP.get(fmt, "720p")
                    else:
                        payload["aspectRatio"] = fmt
                if duration and str(duration) not in ("0", "auto", "none", "null"):
                    try:
                        dur = int(duration)
                        limits = MODEL_DURATION_LIMITS.get(model)
                        if limits:
                            lo, hi = limits
                            if not (lo <= dur <= hi):
                                return JSONResponse(
                                    {"error": f"{model} duration must be {lo}–{hi}s (you sent {dur}s)"},
                                    status_code=400
                                )
                        payload["duration"] = dur
                    except (ValueError, TypeError):
                        pass

                logging.info("vs_generate final payload: %s", json.dumps(payload))
                result = await arcads_generate_video(payload)
                arcads_id = result.get("id") or result.get("videoId", "")

                conn = sqlite3.connect(DB_PATH)
                conn.execute("""
                    INSERT INTO video_jobs (id, client, job_type, model, prompt, status, arcads_id, formats)
                    VALUES (?,?,?,?,?,?,?,?)
                """, (vid_job_id, client_val, "video", model, prompt, "pending", arcads_id, json.dumps({"format": fmt})))
                conn.commit()
                conn.close()
                jobs_created.append({"jobId": vid_job_id, "arcadsId": arcads_id, "format": fmt})

        return {"success": True, "jobs": jobs_created, "sessionId": job_id}
    except Exception as e:
        logging.error("Video generate error: %s", traceback.format_exc())
        return JSONResponse({"error": str(e)}, status_code=500)

@app.get("/api/video-studio/status/{arcads_id}")
async def vs_status(arcads_id: str):
    try:
        data = await arcads_poll_video(arcads_id)
        # /v1/assets returns: status ("created"|"pending"|"generated"|"failed"), url
        status = data.get("status") or data.get("videoStatus") or "pending"
        url    = data.get("url") or data.get("videoUrl") or ""
        logging.info("Poll %s → status=%s url=%s", arcads_id, status, url[:60] if url else "")
        return {"arcadsId": arcads_id, "status": status, "url": url, "raw": data}
    except Exception as e:
        return JSONResponse({"error": str(e), "status": "error"}, status_code=200)

@app.post("/api/video-studio/mimic")
async def vs_mimic(
    referenceFile: UploadFile = File(...),
    prompt: str = Form(""),
    model: str = Form("kling-3.0"),
    productId: str = Form(""),
    client: str = Form("other"),
    formats: str = Form('["9:16"]'),
    variations: int = Form(1),
):
    try:
        if not ARCADS_CLIENT_ID:
            return JSONResponse({"error": "Arcads credentials not configured"}, status_code=400)

        file_bytes    = await referenceFile.read()
        content_type  = referenceFile.content_type or "video/mp4"
        is_video      = content_type.startswith("video/")
        file_path     = await arcads_upload_file(file_bytes, referenceFile.filename, content_type)
        logging.info("Mimic upload complete: filePath=%s model=%s is_video=%s", file_path, model, is_video)

        # Auto-resolve productId
        product_id = (productId or "").strip()
        if not product_id:
            try:
                products = await arcads_get_products()
                if products:
                    product_id = products[0].get("id", "")
            except Exception:
                pass

        fmt_list = json.loads(formats) if isinstance(formats, str) else formats
        jobs_created = []

        for fmt in fmt_list:
            for _ in range(variations):
                vid_job_id = str(uuid.uuid4())
                payload = {
                    "model": model,
                    "prompt": prompt or "Recreate the visual style, mood, and composition of this reference",
                }
                if model == "seedance-2.0":
                    payload["resolution"] = SEEDANCE2_RESOLUTION_MAP.get(fmt, "720p")
                else:
                    payload["aspectRatio"] = fmt
                if product_id:
                    payload["productId"] = product_id

                # Route reference file to the correct field based on model + file type
                STARTFRAME_MODELS = {"kling-3.0", "kling-2.6", "veo31", "grok-video"}
                REFIMAGE_MODELS   = {"sora2", "sora2-pro", "seedance"}
                if is_video and model == "seedance-2.0":
                    payload["referenceVideos"] = [file_path]
                elif model in STARTFRAME_MODELS:
                    payload["startFrame"] = file_path
                elif model in REFIMAGE_MODELS or model == "seedance-2.0":
                    payload["referenceImages"] = [file_path]
                else:
                    payload["startFrame"] = file_path

                logging.info("Mimic payload: %s", json.dumps(payload))
                result = await arcads_generate_video(payload)
                arcads_id = result.get("id") or result.get("videoId", "")

                conn = sqlite3.connect(DB_PATH)
                conn.execute("""
                    INSERT INTO video_jobs (id, client, job_type, model, prompt, status, arcads_id, formats)
                    VALUES (?,?,?,?,?,?,?,?)
                """, (vid_job_id, client, "mimic", model, prompt, "pending", arcads_id, json.dumps({"format": fmt})))
                conn.commit()
                conn.close()
                jobs_created.append({"jobId": vid_job_id, "arcadsId": arcads_id, "format": fmt})

        return {"success": True, "jobs": jobs_created, "filePath": file_path}
    except Exception as e:
        logging.error("Mimic error: %s", traceback.format_exc())
        return JSONResponse({"error": str(e)}, status_code=500)

@app.post("/api/video-studio/upload-file")
async def vs_upload_file(file: UploadFile = File(...)):
    """Upload any file to Arcads presigned storage, return filePath."""
    try:
        file_bytes   = await file.read()
        content_type = file.content_type or "application/octet-stream"
        file_path    = await arcads_upload_file(file_bytes, file.filename, content_type)
        return {"filePath": file_path}
    except Exception as e:
        logging.error("Upload file error: %s", traceback.format_exc())
        return JSONResponse({"error": str(e)}, status_code=500)

@app.post("/api/video-studio/brand-visual")
async def vs_brand_visual(request: Request):
    try:
        body = await request.json()
        if not ARCADS_CLIENT_ID:
            return JSONResponse({"error": "Arcads credentials not configured"}, status_code=400)

        product_id = (body.get("productId") or "").strip()
        if not product_id:
            try:
                products = await arcads_get_products()
                if products:
                    product_id = products[0].get("id", "")
                    logging.info("Brand visual auto-resolved productId: %s", product_id)
            except Exception as pe:
                logging.warning("Brand visual product auto-resolve failed: %s", pe)

        payload = {
            "model": body.get("model", "nano-banana-2"),
            "prompt": body.get("prompt", ""),
            "aspectRatio": body.get("aspectRatio", "1:1"),
        }
        if product_id:
            payload["productId"] = product_id
        # referenceImages takes presigned file paths — if provided, pass as array
        if body.get("referenceImagePath"):
            payload["referenceImages"] = [body["referenceImagePath"]]

        result = await arcads_generate_image(payload)
        asset_id = result.get("id") or result.get("assetId", "")
        return {"success": True, "assetId": asset_id, "raw": result}
    except Exception as e:
        logging.error("Brand visual error: %s", traceback.format_exc())
        return JSONResponse({"error": str(e)}, status_code=500)

@app.get("/api/video-studio/asset/{asset_id}")
async def vs_asset_status(asset_id: str):
    try:
        data = await arcads_poll_asset(asset_id)
        return data
    except Exception as e:
        return JSONResponse({"error": str(e), "status": "error"}, status_code=200)

@app.get("/api/video-studio/jobs")
async def vs_list_jobs(client: str = "all"):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    if client == "all":
        rows = conn.execute("SELECT * FROM video_jobs ORDER BY created_at DESC LIMIT 50").fetchall()
    else:
        rows = conn.execute("SELECT * FROM video_jobs WHERE client=? ORDER BY created_at DESC LIMIT 50", (client,)).fetchall()
    conn.close()
    return {"jobs": [dict(r) for r in rows]}

@app.post("/api/video-studio/save-result")
async def vs_save_result(request: Request):
    """Called by frontend when polling completes — saves result URL to DB."""
    body = await request.json()
    arcads_id  = body.get("arcadsId", "")
    status     = body.get("status", "")
    result_url = body.get("url", "")
    if arcads_id:
        conn = sqlite3.connect(DB_PATH)
        conn.execute(
            "UPDATE video_jobs SET status=?, result_url=? WHERE arcads_id=?",
            (status, result_url, arcads_id)
        )
        conn.commit()
        conn.close()
    return {"ok": True}

async def _bg_poll_pending():
    """Background task: poll Arcads for pending video/image jobs and update DB."""
    while True:
        await asyncio.sleep(30)
        try:
            conn = sqlite3.connect(DB_PATH)
            rows = conn.execute(
                "SELECT arcads_id, job_type FROM video_jobs WHERE status='pending' AND arcads_id != '' LIMIT 20"
            ).fetchall()
            conn.close()
            for (arcads_id, job_type) in rows:
                try:
                    # Both video and image jobs are polled via /v1/assets/{id}
                    data   = await arcads_poll_video(arcads_id)
                    status = (data.get("status") or data.get("videoStatus") or "pending").lower()
                    url    = data.get("url") or data.get("videoUrl") or ""
                    if status in ("done", "generated", "completed", "failed", "error"):
                        conn = sqlite3.connect(DB_PATH)
                        conn.execute(
                            "UPDATE video_jobs SET status=?, result_url=? WHERE arcads_id=?",
                            (status, url, arcads_id)
                        )
                        conn.commit()
                        conn.close()
                        logging.info("BG poll: %s [%s] → %s  url=%s", arcads_id, job_type, status, url[:80] if url else "")
                except Exception as e:
                    logging.warning("BG poll error for %s: %s", arcads_id, e)
        except Exception as e:
            logging.warning("BG poll loop error: %s", e)

@app.on_event("startup")
async def startup_event():
    asyncio.create_task(_bg_poll_pending())


if __name__ == "__main__":
    port = int(os.environ.get("PORT", "5050"))
    uvicorn.run(app, host="0.0.0.0", port=port)
